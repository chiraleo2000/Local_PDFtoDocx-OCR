"""Tests for the absolute-position (layout-faithful) DOCX/HTML exporter."""
import base64
from io import BytesIO

import pytest

from src.pipeline import (ContentBlock, _segments_to_lines, _quad_to_rect,
                          _embedded_text_reliable)
from src.exporter import DocumentExporter

docx = pytest.importorskip("docx")
from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from PIL import Image  # noqa: E402


def _png_b64(w=200, h=120):
    img = Image.new("RGB", (w, h), (200, 60, 60))
    buf = BytesIO()
    img.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode()


PW, PH = 1240, 1754  # page units (px @ ~150 dpi A4)


@pytest.fixture
def blocks():
    return [
        ContentBlock(
            "text", 0, 100, 200,
            text="หัวข้อเอกสารทดสอบ\nบรรทัดที่หนึ่ง ของเนื้อหา",
            bbox=[200, 100, 1040, 260], page_width=PW, page_height=PH,
            lines=[
                {"text": "หัวข้อเอกสารทดสอบ", "bbox": [450, 100, 790, 140]},
                {"text": "บรรทัดที่หนึ่ง ของเนื้อหา", "bbox": [200, 160, 900, 195]},
            ]),
        ContentBlock(
            "table", 0, 400, 150,
            text="A\tB",
            table_html=("<table><tr><th>ชื่อ</th><th>จำนวน</th></tr>"
                        "<tr><td>ข้าว</td><td>10</td></tr></table>"),
            bbox=[150, 400, 1090, 700], page_width=PW, page_height=PH,
            table_meta={"col_widths": [0.7, 0.3]}),
        ContentBlock(
            "figure", 0, 800, 300,
            figure={"base64": _png_b64(), "index": 1, "page": 0,
                    "width": 200, "height": 120},
            bbox=[300, 800, 940, 1180], page_width=PW, page_height=PH),
    ]


def test_absolute_docx(blocks):
    exp = DocumentExporter()
    files = exp.create_all_from_blocks(blocks, "Pages: 1", "A4", "Normal",
                                       layout_mode="absolute")
    assert files["docx"]
    doc = Document(files["docx"])
    xml = doc.element.xml

    # Text placed in page-anchored frames at exact coordinates
    assert "framePr" in xml
    assert 'vAnchor="page"' in xml

    # Table is a real editable Word table, absolutely positioned
    assert "tblpPr" in xml
    assert len(doc.tables) == 1
    tbl = doc.tables[0]
    assert tbl.cell(0, 0).text == "ชื่อ"
    assert tbl.cell(1, 1).text == "10"

    # Figure embedded as image
    assert len(doc.inline_shapes) == 1

    # tblPr children must respect OOXML schema order
    order = ["tblStyle", "tblpPr", "tblOverlap", "bidiVisual",
             "tblStyleRowBandSize", "tblStyleColBandSize", "tblW", "jc",
             "tblCellSpacing", "tblInd", "tblBorders", "shd", "tblLayout",
             "tblCellMar", "tblLook", "tblCaption", "tblDescription"]
    rank = {qn(f"w:{t}"): i for i, t in enumerate(order)}
    ranks = [rank.get(el.tag, 99) for el in tbl._tbl.tblPr]
    assert ranks == sorted(ranks)


def test_absolute_html(blocks):
    exp = DocumentExporter()
    files = exp.create_all_from_blocks(blocks, layout_mode="absolute")
    with open(files["html"], encoding="utf-8") as f:
        html_out = f.read()
    assert "position: absolute" in html_out
    assert "class='page'" in html_out
    assert "data:image/png;base64" in html_out


def test_flow_fallback_without_positions():
    exp = DocumentExporter()
    flow_blocks = [ContentBlock("text", 0, 0, 0, text="plain text only")]
    files = exp.create_all_from_blocks(flow_blocks, layout_mode="absolute")
    assert files["docx"]
    Document(files["docx"])  # opens cleanly


def test_segments_to_lines_clustering():
    segs = [
        {"text": "Hello", "bbox": [[10, 10], [60, 10], [60, 30], [10, 30]]},
        {"text": "world", "bbox": [[70, 12], [120, 12], [120, 32], [70, 32]]},
        {"text": "next line", "bbox": [[10, 60], [120, 60], [120, 80], [10, 80]]},
    ]
    lines = _segments_to_lines(segs)
    assert len(lines) == 2
    assert lines[0]["text"] == "Hello world"
    assert lines[0]["bbox"] == [10, 10, 120, 32]


def test_quad_to_rect():
    assert _quad_to_rect([[10, 20], [110, 20], [110, 50], [10, 50]]) == \
        [10, 20, 110, 50]
    assert _quad_to_rect([1, 2, 3, 4]) == [1.0, 2.0, 3.0, 4.0]
    assert _quad_to_rect(None) is None


def test_broken_thai_text_layer_rejected():
    """Pages whose embedded text layer lost its Thai glyphs must be OCRed.

    Regression for testocrtor-demo.pdf: the PDF carries a hidden text
    layer in which every Thai character is missing (no ToUnicode map),
    so the old fast path exported ASCII garbage instead of OCRing.
    """
    page0 = ('ICT01 \n 2 email ict_its@opsmoacgoth \n'
             '3)4)  5)  6)  7)  8)  9)  10)  11)   7)  \n ')
    page1 = ('ICT01 \n 3 email ict_its@opsmoacgoth \n2.2.1\n'
             '16  PC   728 166 97 Network) Next Generation Firewall) - '
             'Log File  Access Switch) Access Point) 1 1 89 32 Printer)')
    assert not _embedded_text_reliable(page0, "tha+eng")
    assert not _embedded_text_reliable(page1, "tha+eng")
    # Healthy layers keep the fast path
    assert _embedded_text_reliable(
        "รายงานประจำปี กระทรวงเกษตรและสหกรณ์ สรุปผลการดำเนินงาน", "tha+eng")
    assert _embedded_text_reliable(
        "Annual report covering network and security systems.", "eng")
    # Replacement characters = broken CMap
    assert not _embedded_text_reliable("��� broken ��", "eng")
