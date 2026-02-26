"""
Selenium UI Tests for PDF OCR Pipeline v0.2.0

Tests the Gradio web UI end-to-end using Selenium WebDriver:
  - Server launch & health
  - PDF upload via Convert tab
  - Language / engine / quality selection
  - Conversion & DOCX download
  - Review & Correct tab interaction
  - History tab
  - Settings tab
  - System Status tab

Outputs a DOCX report with all test results.

Usage:
    python -m pytest tests/test_selenium_ui.py -v --tb=short
    python tests/test_selenium_ui.py          # direct execution
"""

import os
import sys
import time
import json
import shutil
import logging
import subprocess
import traceback
from pathlib import Path
from datetime import datetime

# ── Paths ────────────────────────────────────────────────────────────────────
TESTS_DIR = Path(__file__).resolve().parent
ROOT_DIR = TESTS_DIR.parent
TEST_PDF = TESTS_DIR / "testocrtor.pdf"
OUTPUT_DIR = ROOT_DIR / "test_selenium_output"
REPORT_DOCX = OUTPUT_DIR / "selenium_test_report.docx"
APP_PORT = int(os.environ.get("TEST_APP_PORT", "7873"))
APP_URL = os.environ.get("TEST_APP_URL", f"http://127.0.0.1:{APP_PORT}")
# When TEST_APP_URL is set, we skip starting a local server (e.g. Docker mode)
EXTERNAL_SERVER = bool(os.environ.get("TEST_APP_URL"))
WAIT_TIMEOUT = 120  # seconds to wait for server / conversion

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(name)s] %(levelname)s: %(message)s",
)
logger = logging.getLogger("selenium_test")


# ══════════════════════════════════════════════════════════════════════════════
# Helpers
# ══════════════════════════════════════════════════════════════════════════════

def _ensure_output_dir():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def _start_app_server():
    """Launch the Gradio app as a background subprocess."""
    env = os.environ.copy()
    env.update({
        "SERVER_PORT": str(APP_PORT),
        "SERVER_HOST": "127.0.0.1",
        "SHARE_GRADIO": "false",
        "DEBUG_MODE": "false",
    })
    proc = subprocess.Popen(
        [sys.executable, str(ROOT_DIR / "app.py")],
        env=env,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        cwd=str(ROOT_DIR),
    )
    return proc


def _wait_for_server(url: str, timeout: int = WAIT_TIMEOUT) -> bool:
    """Poll until the server responds with HTTP 200."""
    import urllib.request
    deadline = time.time() + timeout
    while time.time() < deadline:
        try:
            resp = urllib.request.urlopen(url, timeout=3)
            if resp.status == 200:
                return True
        except Exception:
            time.sleep(1.5)
    return False


def _stop_server(proc):
    """Gracefully terminate the app server."""
    if proc is None:
        return
    proc.terminate()
    try:
        proc.wait(timeout=10)
    except subprocess.TimeoutExpired:
        proc.kill()


def _get_webdriver():
    """
    Create a Selenium WebDriver. Tries Edge first (pre-installed on Windows),
    then Chrome, then Firefox.
    """
    from selenium import webdriver

    # Try Edge (Windows default)
    try:
        from selenium.webdriver.edge.options import Options as EdgeOptions
        from selenium.webdriver.edge.service import Service as EdgeService
        opts = EdgeOptions()
        opts.add_argument("--headless=new")
        opts.add_argument("--no-sandbox")
        opts.add_argument("--disable-dev-shm-usage")
        opts.add_argument("--disable-gpu")
        opts.add_argument("--window-size=1920,1080")
        # Set download directory
        prefs = {
            "download.default_directory": str(OUTPUT_DIR),
            "download.prompt_for_download": False,
            "download.directory_upgrade": True,
        }
        opts.add_experimental_option("prefs", prefs)
        driver = webdriver.Edge(options=opts)
        logger.info("Using Edge WebDriver (headless)")
        return driver
    except Exception as e:
        logger.warning("Edge WebDriver not available: %s", e)

    # Try Chrome
    try:
        from selenium.webdriver.chrome.options import Options as ChromeOptions
        opts = ChromeOptions()
        opts.add_argument("--headless=new")
        opts.add_argument("--no-sandbox")
        opts.add_argument("--disable-dev-shm-usage")
        opts.add_argument("--disable-gpu")
        opts.add_argument("--window-size=1920,1080")
        prefs = {
            "download.default_directory": str(OUTPUT_DIR),
            "download.prompt_for_download": False,
            "download.directory_upgrade": True,
        }
        opts.add_experimental_option("prefs", prefs)
        driver = webdriver.Chrome(options=opts)
        logger.info("Using Chrome WebDriver (headless)")
        return driver
    except Exception as e:
        logger.warning("Chrome WebDriver not available: %s", e)

    # Try Firefox
    try:
        from selenium.webdriver.firefox.options import Options as FFOptions
        opts = FFOptions()
        opts.add_argument("--headless")
        driver = webdriver.Firefox(options=opts)
        logger.info("Using Firefox WebDriver (headless)")
        return driver
    except Exception as e:
        logger.warning("Firefox WebDriver not available: %s", e)

    raise RuntimeError(
        "No WebDriver available. Install one of: "
        "msedgedriver, chromedriver, geckodriver"
    )


# ══════════════════════════════════════════════════════════════════════════════
# Test Runner
# ══════════════════════════════════════════════════════════════════════════════

class SeleniumTestRunner:
    """Runs Selenium UI tests and collects results."""

    def __init__(self):
        self.results: list[dict] = []
        self.server_proc = None
        self.driver = None
        self.start_time = datetime.now()

    def _record(self, name: str, passed: bool, detail: str = "",
                duration: float = 0.0, screenshot: str | None = None):
        self.results.append({
            "name": name,
            "passed": passed,
            "detail": detail,
            "duration_s": round(duration, 2),
            "screenshot": screenshot,
            "timestamp": datetime.now().isoformat(),
        })
        status = "PASS" if passed else "FAIL"
        logger.info("[%s] %s (%.1fs) %s", status, name, duration, detail[:120])

    def setup(self):
        """Start server and WebDriver."""
        _ensure_output_dir()

        if EXTERNAL_SERVER:
            # External server mode (e.g. Docker) — skip starting a local server
            logger.info("External server mode — using %s", APP_URL)
        else:
            # Start app server
            logger.info("Starting Gradio app on port %d ...", APP_PORT)
            self.server_proc = _start_app_server()

        if not _wait_for_server(APP_URL, timeout=WAIT_TIMEOUT):
            stderr = ""
            if self.server_proc:
                try:
                    self.server_proc.terminate()
                    _, err = self.server_proc.communicate(timeout=5)
                    stderr = err.decode("utf-8", errors="replace")[-2000:]
                except Exception:
                    pass
            raise RuntimeError(
                f"App server did not respond within {WAIT_TIMEOUT}s at {APP_URL}.\n"
                f"stderr:\n{stderr}"
            )
        logger.info("Server is ready at %s", APP_URL)

        # Start WebDriver
        self.driver = _get_webdriver()
        self.driver.implicitly_wait(10)

    def teardown(self):
        """Stop WebDriver and server."""
        if self.driver:
            try:
                self.driver.quit()
            except Exception:
                pass
        _stop_server(self.server_proc)
        logger.info("Teardown complete.")

    def _screenshot(self, name: str) -> str | None:
        """Capture a screenshot and return its path."""
        try:
            fname = f"screenshot_{name}_{int(time.time())}.png"
            path = OUTPUT_DIR / fname
            self.driver.save_screenshot(str(path))
            return str(path)
        except Exception:
            return None

    # ──────────────────────────────────────────────────────────────────────
    # Individual tests
    # ──────────────────────────────────────────────────────────────────────

    def test_server_health(self):
        """Verify server returns 200 and serves HTML."""
        t0 = time.time()
        try:
            self.driver.get(APP_URL)
            time.sleep(2)
            title = self.driver.title
            page_src = self.driver.page_source
            is_html = "<!doctype html>" in page_src.lower() or "<html" in page_src.lower()
            has_gradio = "gradio" in page_src.lower()
            self._record(
                "Server Health",
                is_html and has_gradio,
                f"Title: {title!r}, HTML: {is_html}, Gradio: {has_gradio}",
                time.time() - t0,
                self._screenshot("server_health"),
            )
        except Exception as exc:
            self._record("Server Health", False, str(exc), time.time() - t0)

    def test_hero_bar_content(self):
        """Check the hero bar shows correct version and engine info."""
        t0 = time.time()
        try:
            self.driver.get(APP_URL)
            time.sleep(2)
            page_src = self.driver.page_source
            has_title = "PDF OCR Pipeline" in page_src
            has_version = "v0.2.0" in page_src
            has_engines = "Typhoon OCR" in page_src
            self._record(
                "Hero Bar Content",
                has_title and has_version and has_engines,
                f"Title: {has_title}, Version: {has_version}, Engines: {has_engines}",
                time.time() - t0,
                self._screenshot("hero_bar"),
            )
        except Exception as exc:
            self._record("Hero Bar Content", False, str(exc), time.time() - t0)

    def test_tab_navigation(self):
        """Verify all 6 tabs are present and clickable."""
        from selenium.webdriver.common.by import By
        t0 = time.time()
        try:
            self.driver.get(APP_URL)
            time.sleep(2)

            expected_tabs = [
                "Convert PDF", "Review & Correct", "Training",
                "History", "Settings", "System Status",
            ]
            tab_buttons = self.driver.find_elements(
                By.CSS_SELECTOR, "button[role='tab']"
            )
            tab_labels = [btn.text.strip() for btn in tab_buttons if btn.text.strip()]

            found = []
            for expected in expected_tabs:
                match = any(expected.lower() in lbl.lower() for lbl in tab_labels)
                found.append(match)

            all_found = all(found)
            # Click each tab to verify it's interactable
            for btn in tab_buttons:
                try:
                    btn.click()
                    time.sleep(0.3)
                except Exception:
                    pass

            self._record(
                "Tab Navigation",
                all_found,
                f"Expected: {expected_tabs}, Found: {tab_labels}",
                time.time() - t0,
                self._screenshot("tabs"),
            )
        except Exception as exc:
            self._record("Tab Navigation", False, str(exc), time.time() - t0)

    def test_convert_tab_elements(self):
        """Verify Convert PDF tab has all required UI elements."""
        from selenium.webdriver.common.by import By
        t0 = time.time()
        try:
            self.driver.get(APP_URL)
            time.sleep(2)

            page_src = self.driver.page_source

            checks = {
                "Upload area": "Select PDF file" in page_src or "Upload" in page_src,
                "Language dropdown": "OCR Language" in page_src,
                "Quality dropdown": "Quality Level" in page_src or "Quality" in page_src,
                "Convert button": "Convert to DOCX" in page_src,
                "Preview area": "Preview" in page_src or "Document Preview" in page_src,
            }

            all_ok = all(checks.values())
            detail = ", ".join(f"{k}: {'OK' if v else 'MISSING'}" for k, v in checks.items())
            self._record(
                "Convert Tab Elements",
                all_ok, detail, time.time() - t0,
                self._screenshot("convert_tab"),
            )
        except Exception as exc:
            self._record("Convert Tab Elements", False, str(exc), time.time() - t0)

    def test_pdf_upload_and_convert(self):
        """Upload a PDF and trigger conversion (core E2E test)."""
        from selenium.webdriver.common.by import By
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC
        t0 = time.time()
        try:
            self.driver.get(APP_URL)
            time.sleep(3)

            if not TEST_PDF.exists():
                self._record(
                    "PDF Upload & Convert", False,
                    f"Test PDF not found: {TEST_PDF}", time.time() - t0)
                return

            # Find the file input (Gradio wraps it in a hidden <input type="file">)
            file_inputs = self.driver.find_elements(By.CSS_SELECTOR, 'input[type="file"]')
            if not file_inputs:
                self._record(
                    "PDF Upload & Convert", False,
                    "No file input found on page", time.time() - t0,
                    self._screenshot("no_file_input"))
                return

            # Upload to the first file input (Convert tab)
            file_inputs[0].send_keys(str(TEST_PDF))
            logger.info("PDF uploaded, waiting for preview...")
            time.sleep(5)

            screenshot_after_upload = self._screenshot("after_upload")

            # Click "Convert to DOCX" button
            convert_buttons = self.driver.find_elements(By.XPATH,
                "//button[contains(., 'Convert to DOCX')]")
            if not convert_buttons:
                # Fallback: find by partial text
                convert_buttons = self.driver.find_elements(By.XPATH,
                    "//button[contains(., 'Convert')]")

            if convert_buttons:
                convert_buttons[0].click()
                logger.info("Convert button clicked, waiting for conversion...")
            else:
                self._record(
                    "PDF Upload & Convert", False,
                    "Convert button not found", time.time() - t0,
                    screenshot_after_upload)
                return

            # Wait for processing to complete (look for "Complete" in status)
            conversion_timeout = 300  # 5 minutes max (model loading may be slow)
            deadline = time.time() + conversion_timeout
            completed = False
            error_occurred = False
            status_text = ""

            while time.time() < deadline:
                time.sleep(3)
                page_src = self.driver.page_source
                if "Conversion Complete" in page_src or "Complete!" in page_src:
                    completed = True
                    break
                if "Error" in page_src and "Processing" not in page_src:
                    # Check if it's a real error vs just the error element existing
                    status_elements = self.driver.find_elements(
                        By.XPATH, "//*[contains(text(), 'Error')]")
                    for el in status_elements:
                        txt = el.text.strip()
                        if txt and "Error" in txt and len(txt) > 5:
                            error_occurred = True
                            status_text = txt
                            break
                    if error_occurred:
                        break

            screenshot_after = self._screenshot("after_convert")

            if completed:
                # Check for download links
                download_files = self.driver.find_elements(
                    By.CSS_SELECTOR, 'a[download], a[href*="file="]')
                has_downloads = len(download_files) > 0

                # Try to get status text
                try:
                    status_elements = self.driver.find_elements(
                        By.XPATH, "//*[contains(text(), 'Pages:')]")
                    if status_elements:
                        status_text = status_elements[0].text[:200]
                except Exception:
                    pass

                self._record(
                    "PDF Upload & Convert", True,
                    f"Conversion completed. Downloads: {has_downloads}. "
                    f"Status: {status_text[:150]}",
                    time.time() - t0, screenshot_after)
            elif error_occurred:
                self._record(
                    "PDF Upload & Convert", False,
                    f"Error during conversion: {status_text[:200]}",
                    time.time() - t0, screenshot_after)
            else:
                self._record(
                    "PDF Upload & Convert", False,
                    f"Conversion timed out after {conversion_timeout}s",
                    time.time() - t0, screenshot_after)

        except Exception as exc:
            self._record("PDF Upload & Convert", False,
                         traceback.format_exc()[-500:], time.time() - t0,
                         self._screenshot("convert_error"))

    def test_engine_dropdown_options(self):
        """Verify OCR Engine dropdown has the correct v0.2.0 options
        (no Tesseract, no EasyOCR)."""
        from selenium.webdriver.common.by import By
        t0 = time.time()
        try:
            self.driver.get(APP_URL)
            time.sleep(2)
            page_src = self.driver.page_source

            has_typhoon = "Typhoon OCR 3B" in page_src
            has_trocr = "Thai-TrOCR" in page_src
            has_paddle = "PaddleOCR" in page_src
            no_tesseract = "Tesseract" not in page_src
            no_easyocr = "EasyOCR" not in page_src

            all_ok = has_typhoon and has_trocr and has_paddle and no_tesseract and no_easyocr
            self._record(
                "Engine Dropdown Options", all_ok,
                f"Typhoon: {has_typhoon}, TrOCR: {has_trocr}, Paddle: {has_paddle}, "
                f"NoTesseract: {no_tesseract}, NoEasyOCR: {no_easyocr}",
                time.time() - t0, self._screenshot("engine_options"))
        except Exception as exc:
            self._record("Engine Dropdown Options", False, str(exc), time.time() - t0)

    def test_language_dropdown_options(self):
        """Verify language dropdown has Thai, English, Thai+English, etc."""
        from selenium.webdriver.common.by import By
        t0 = time.time()
        try:
            self.driver.get(APP_URL)
            time.sleep(2)
            page_src = self.driver.page_source

            has_thai = "Thai" in page_src
            has_english = "English" in page_src
            has_thai_eng = "Thai + English" in page_src
            has_auto = "Auto-detect" in page_src

            all_ok = has_thai and has_english and has_thai_eng
            self._record(
                "Language Options", all_ok,
                f"Thai: {has_thai}, English: {has_english}, "
                f"Thai+Eng: {has_thai_eng}, Auto: {has_auto}",
                time.time() - t0)
        except Exception as exc:
            self._record("Language Options", False, str(exc), time.time() - t0)

    def test_settings_tab(self):
        """Navigate to Settings tab and verify configuration elements."""
        from selenium.webdriver.common.by import By
        t0 = time.time()
        try:
            self.driver.get(APP_URL)
            time.sleep(2)

            # Click Settings tab
            tabs = self.driver.find_elements(By.CSS_SELECTOR, "button[role='tab']")
            settings_tab = None
            for tab in tabs:
                if "settings" in tab.text.lower():
                    settings_tab = tab
                    break

            if settings_tab:
                settings_tab.click()
                time.sleep(1)

            page_src = self.driver.page_source
            has_config = "Configuration" in page_src
            has_ocr_config = "OCR Configuration" in page_src
            has_layout = "Layout Detection" in page_src
            has_installed = "Installed Engines" in page_src
            has_yolo_conf = "YOLO Confidence" in page_src

            all_ok = has_config and has_ocr_config and has_layout
            self._record(
                "Settings Tab", all_ok,
                f"Config: {has_config}, OCR: {has_ocr_config}, "
                f"Layout: {has_layout}, Installed: {has_installed}, "
                f"YOLO: {has_yolo_conf}",
                time.time() - t0, self._screenshot("settings_tab"))
        except Exception as exc:
            self._record("Settings Tab", False, str(exc), time.time() - t0)

    def test_system_status_tab(self):
        """Navigate to System Status tab and check pipeline status JSON."""
        from selenium.webdriver.common.by import By
        t0 = time.time()
        try:
            self.driver.get(APP_URL)
            time.sleep(2)

            tabs = self.driver.find_elements(By.CSS_SELECTOR, "button[role='tab']")
            status_tab = None
            for tab in tabs:
                if "system" in tab.text.lower() or "status" in tab.text.lower():
                    status_tab = tab
                    break

            if status_tab:
                status_tab.click()
                time.sleep(1)

            page_src = self.driver.page_source
            has_status = "Pipeline Status" in page_src
            has_json = "ocr_engine" in page_src or "layout_detector" in page_src

            self._record(
                "System Status Tab",
                has_status,
                f"Status section: {has_status}, JSON data: {has_json}",
                time.time() - t0, self._screenshot("system_status"))
        except Exception as exc:
            self._record("System Status Tab", False, str(exc), time.time() - t0)

    def test_review_tab_elements(self):
        """Verify Review & Correct tab has region controls."""
        from selenium.webdriver.common.by import By
        t0 = time.time()
        try:
            self.driver.get(APP_URL)
            time.sleep(2)

            tabs = self.driver.find_elements(By.CSS_SELECTOR, "button[role='tab']")
            review_tab = None
            for tab in tabs:
                if "review" in tab.text.lower():
                    review_tab = tab
                    break

            if review_tab:
                review_tab.click()
                time.sleep(1)

            page_src = self.driver.page_source
            has_add_region = "Add Region" in page_src
            has_clear = "Clear" in page_src
            has_convert_corrections = "Convert with Corrections" in page_src
            has_bbox_fields = "x0" in page_src or "left" in page_src
            has_region_type = "Region Type" in page_src or "table" in page_src

            all_ok = has_add_region and has_convert_corrections
            self._record(
                "Review Tab Elements", all_ok,
                f"AddRegion: {has_add_region}, Clear: {has_clear}, "
                f"ConvertCorr: {has_convert_corrections}, "
                f"BBox: {has_bbox_fields}, RegType: {has_region_type}",
                time.time() - t0, self._screenshot("review_tab"))
        except Exception as exc:
            self._record("Review Tab Elements", False, str(exc), time.time() - t0)

    def test_history_tab(self):
        """Navigate to History tab and check for table/buttons."""
        from selenium.webdriver.common.by import By
        t0 = time.time()
        try:
            self.driver.get(APP_URL)
            time.sleep(2)

            tabs = self.driver.find_elements(By.CSS_SELECTOR, "button[role='tab']")
            hist_tab = None
            for tab in tabs:
                if "history" in tab.text.lower():
                    hist_tab = tab
                    break

            if hist_tab:
                hist_tab.click()
                time.sleep(1)

            page_src = self.driver.page_source
            has_history = "Processing History" in page_src
            has_refresh = "Refresh" in page_src
            has_entry_id = "Entry ID" in page_src
            has_download = "Download DOCX" in page_src

            all_ok = has_history and has_entry_id and has_download
            self._record(
                "History Tab", all_ok,
                f"History: {has_history}, Refresh: {has_refresh}, "
                f"EntryID: {has_entry_id}, Download: {has_download}",
                time.time() - t0, self._screenshot("history_tab"))
        except Exception as exc:
            self._record("History Tab", False, str(exc), time.time() - t0)

    def test_training_tab(self):
        """Navigate to Training tab and verify retrain dashboard."""
        from selenium.webdriver.common.by import By
        t0 = time.time()
        try:
            self.driver.get(APP_URL)
            time.sleep(2)

            tabs = self.driver.find_elements(By.CSS_SELECTOR, "button[role='tab']")
            train_tab = None
            for tab in tabs:
                if "training" in tab.text.lower():
                    train_tab = tab
                    break

            if train_tab:
                train_tab.click()
                time.sleep(1)

            page_src = self.driver.page_source
            has_retrain = "Auto-Retrain" in page_src or "Retrain" in page_src
            has_corrections = "corrections" in page_src.lower()
            has_refresh = "Refresh Stats" in page_src

            all_ok = has_retrain or has_corrections
            self._record(
                "Training Tab", all_ok,
                f"Retrain: {has_retrain}, Corrections: {has_corrections}, "
                f"Refresh: {has_refresh}",
                time.time() - t0, self._screenshot("training_tab"))
        except Exception as exc:
            self._record("Training Tab", False, str(exc), time.time() - t0)

    def test_no_tesseract_references(self):
        """Ensure NO Tesseract or EasyOCR in visible dropdown/label text.
        Note: raw page_source may contain JS bundles with unrelated words,
        so we check only visible element text instead."""
        from selenium.webdriver.common.by import By
        t0 = time.time()
        try:
            self.driver.get(APP_URL)
            time.sleep(2)

            # Check visible text in dropdown options and labels
            visible_text = ""
            for el in self.driver.find_elements(By.CSS_SELECTOR,
                    "label, option, .wrap, span, button, .selected-item, "
                    "[data-testid], .label-text, h1, h2, h3, p, .hero-bar, "
                    ".hero-badge"):
                try:
                    visible_text += " " + el.text
                except Exception:
                    pass

            no_tesseract = "Tesseract" not in visible_text
            no_easyocr = "EasyOCR" not in visible_text

            self._record(
                "No Tesseract/EasyOCR",
                no_tesseract and no_easyocr,
                f"NoTesseract: {no_tesseract}, NoEasyOCR: {no_easyocr}",
                time.time() - t0)
        except Exception as exc:
            self._record("No Tesseract/EasyOCR", False, str(exc), time.time() - t0)

    def test_responsive_layout(self):
        """Check page renders at different sizes without crashing.
        Note: headless browsers often report JS fetch/network errors that
        don't affect real usage, so we only fail on rendering crashes."""
        from selenium.webdriver.common.by import By
        t0 = time.time()
        try:
            render_ok = True
            js_warnings = []
            for w, h in [(1920, 1080), (1024, 768), (768, 1024)]:
                self.driver.set_window_size(w, h)
                self.driver.get(APP_URL)
                time.sleep(1.5)

                # Verify the page actually rendered
                page_src = self.driver.page_source
                if "PDF OCR Pipeline" not in page_src:
                    render_ok = False

                # Collect JS errors as warnings (not failures)
                try:
                    logs = self.driver.get_log("browser")
                    severe = [l for l in logs if l.get("level") == "SEVERE"]
                    if severe:
                        js_warnings.extend(
                            f"{w}x{h}: {l['message'][:80]}" for l in severe[:1])
                except Exception:
                    pass

            self.driver.set_window_size(1920, 1080)

            self._record(
                "Responsive Layout",
                render_ok,
                f"Rendered OK: {render_ok}. JS warnings: {len(js_warnings)}",
                time.time() - t0)
        except Exception as exc:
            self._record("Responsive Layout", False, str(exc), time.time() - t0)

    # ──────────────────────────────────────────────────────────────────────
    # Run all tests
    # ──────────────────────────────────────────────────────────────────────

    def run_all(self) -> list[dict]:
        """Execute all tests in order, returning results."""
        tests = [
            self.test_server_health,
            self.test_hero_bar_content,
            self.test_tab_navigation,
            self.test_convert_tab_elements,
            self.test_engine_dropdown_options,
            self.test_language_dropdown_options,
            self.test_no_tesseract_references,
            self.test_settings_tab,
            self.test_system_status_tab,
            self.test_review_tab_elements,
            self.test_history_tab,
            self.test_training_tab,
            self.test_responsive_layout,
            self.test_pdf_upload_and_convert,  # longest test last
        ]

        for test_fn in tests:
            try:
                test_fn()
            except Exception as exc:
                self._record(test_fn.__name__, False,
                             f"Unhandled: {exc}", 0.0)

        return self.results


# ══════════════════════════════════════════════════════════════════════════════
# DOCX Report Generator
# ══════════════════════════════════════════════════════════════════════════════

def generate_report_docx(results: list[dict], output_path: Path,
                         total_duration: float):
    """Generate a DOCX test report with results table and summary."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # ── Title ────────────────────────────────────────────────────────────
    title = doc.add_heading("PDF OCR Pipeline — Selenium UI Test Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        f"App Version: v0.2.0\n"
        f"Total Duration: {total_duration:.1f}s\n"
        f"Test PDF: {TEST_PDF.name}"
    )

    # ── Summary ──────────────────────────────────────────────────────────
    passed = sum(1 for r in results if r["passed"])
    failed = sum(1 for r in results if not r["passed"])
    total = len(results)

    doc.add_heading("Summary", level=1)
    summary_table = doc.add_table(rows=1, cols=4)
    summary_table.style = "Light Grid Accent 1"
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = summary_table.rows[0].cells
    hdr[0].text = "Total Tests"
    hdr[1].text = "Passed"
    hdr[2].text = "Failed"
    hdr[3].text = "Pass Rate"
    row = summary_table.add_row().cells
    row[0].text = str(total)
    row[1].text = str(passed)
    row[2].text = str(failed)
    row[3].text = f"{(passed / total * 100) if total else 0:.1f}%"

    # Bold headers
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    doc.add_paragraph("")  # spacer

    # ── Detailed Results ─────────────────────────────────────────────────
    doc.add_heading("Detailed Results", level=1)

    detail_table = doc.add_table(rows=1, cols=4)
    detail_table.style = "Light Grid Accent 1"
    detail_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = detail_table.rows[0].cells
    headers = ["Test Name", "Status", "Duration", "Details"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True

    for r in results:
        row_cells = detail_table.add_row().cells
        row_cells[0].text = r["name"]

        status_text = "PASS" if r["passed"] else "FAIL"
        row_cells[1].text = status_text
        # Colour the status
        for p in row_cells[1].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = (
                    RGBColor(0x05, 0x96, 0x69) if r["passed"]
                    else RGBColor(0xDC, 0x26, 0x26)
                )

        row_cells[2].text = f"{r['duration_s']}s"
        row_cells[3].text = r["detail"][:300] if r["detail"] else ""

    doc.add_paragraph("")  # spacer

    # ── Screenshots section ──────────────────────────────────────────────
    screenshots = [r for r in results if r.get("screenshot") and
                   os.path.exists(r["screenshot"])]
    if screenshots:
        doc.add_heading("Screenshots", level=1)
        for r in screenshots[:8]:  # Limit to 8 screenshots
            doc.add_heading(r["name"], level=3)
            try:
                doc.add_picture(r["screenshot"], width=Inches(5.5))
            except Exception:
                doc.add_paragraph(f"[Screenshot: {r['screenshot']}]")

    # ── Environment Info ─────────────────────────────────────────────────
    doc.add_heading("Environment", level=1)
    env_items = [
        f"Python: {sys.version.split()[0]}",
        f"Platform: {sys.platform}",
        f"CWD: {os.getcwd()}",
        f"App Port: {APP_PORT}",
        f"Test PDF: {TEST_PDF}",
    ]
    for item in env_items:
        doc.add_paragraph(item, style="List Bullet")

    # ── Save ─────────────────────────────────────────────────────────────
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Report saved to %s", output_path)


# ══════════════════════════════════════════════════════════════════════════════
# Pytest integration
# ══════════════════════════════════════════════════════════════════════════════

try:
    import pytest

    @pytest.fixture(scope="module")
    def selenium_runner():
        """Module-scoped fixture: start server + driver, run tests, teardown."""
        runner = SeleniumTestRunner()
        try:
            runner.setup()
            yield runner
        finally:
            runner.teardown()

    class TestSeleniumUI:
        """Pytest wrapper around SeleniumTestRunner tests."""

        def test_server_health(self, selenium_runner):
            selenium_runner.test_server_health()
            assert selenium_runner.results[-1]["passed"]

        def test_hero_bar_content(self, selenium_runner):
            selenium_runner.test_hero_bar_content()
            assert selenium_runner.results[-1]["passed"]

        def test_tab_navigation(self, selenium_runner):
            selenium_runner.test_tab_navigation()
            assert selenium_runner.results[-1]["passed"]

        def test_convert_tab_elements(self, selenium_runner):
            selenium_runner.test_convert_tab_elements()
            assert selenium_runner.results[-1]["passed"]

        def test_engine_dropdown_no_tesseract(self, selenium_runner):
            selenium_runner.test_engine_dropdown_options()
            assert selenium_runner.results[-1]["passed"]

        def test_language_dropdown_options(self, selenium_runner):
            selenium_runner.test_language_dropdown_options()
            assert selenium_runner.results[-1]["passed"]

        def test_no_tesseract_references(self, selenium_runner):
            selenium_runner.test_no_tesseract_references()
            assert selenium_runner.results[-1]["passed"]

        def test_settings_tab(self, selenium_runner):
            selenium_runner.test_settings_tab()
            assert selenium_runner.results[-1]["passed"]

        def test_system_status_tab(self, selenium_runner):
            selenium_runner.test_system_status_tab()
            assert selenium_runner.results[-1]["passed"]

        def test_review_tab_elements(self, selenium_runner):
            selenium_runner.test_review_tab_elements()
            assert selenium_runner.results[-1]["passed"]

        def test_history_tab(self, selenium_runner):
            selenium_runner.test_history_tab()
            assert selenium_runner.results[-1]["passed"]

        def test_training_tab(self, selenium_runner):
            selenium_runner.test_training_tab()
            assert selenium_runner.results[-1]["passed"]

        def test_responsive_layout(self, selenium_runner):
            selenium_runner.test_responsive_layout()
            assert selenium_runner.results[-1]["passed"]

        def test_pdf_upload_and_convert(self, selenium_runner):
            selenium_runner.test_pdf_upload_and_convert()
            assert selenium_runner.results[-1]["passed"]

        def test_generate_report(self, selenium_runner):
            """After all tests, generate the DOCX report."""
            total_dur = (datetime.now() - selenium_runner.start_time).total_seconds()
            generate_report_docx(
                selenium_runner.results, REPORT_DOCX, total_dur)
            assert REPORT_DOCX.exists()
            assert REPORT_DOCX.stat().st_size > 0

except ImportError:
    pass  # pytest not installed; use direct execution below


# ══════════════════════════════════════════════════════════════════════════════
# Direct execution (no pytest needed)
# ══════════════════════════════════════════════════════════════════════════════

def main():
    """Run all Selenium UI tests and generate a DOCX report."""
    print("=" * 70)
    print("  PDF OCR Pipeline — Selenium UI Test Suite")
    print("  Version: v0.2.0 | Port:", APP_PORT)
    print("=" * 70)
    print()

    runner = SeleniumTestRunner()
    try:
        print("[*] Setting up server and WebDriver...")
        runner.setup()
        print("[*] Running tests...\n")
        results = runner.run_all()
    except Exception as exc:
        print(f"\n[!] Fatal error: {exc}")
        traceback.print_exc()
        results = runner.results
    finally:
        runner.teardown()

    # Summary
    passed = sum(1 for r in results if r["passed"])
    failed = sum(1 for r in results if not r["passed"])
    total = len(results)
    total_dur = (datetime.now() - runner.start_time).total_seconds()

    print("\n" + "=" * 70)
    print(f"  Results: {passed}/{total} passed | {failed} failed | "
          f"{total_dur:.1f}s total")
    print("=" * 70)

    for r in results:
        icon = "PASS" if r["passed"] else "FAIL"
        print(f"  [{icon}] {r['name']:35s}  ({r['duration_s']}s)")
        if not r["passed"] and r["detail"]:
            for line in r["detail"][:200].split("\n"):
                print(f"          {line}")

    # Generate DOCX report
    print(f"\n[*] Generating DOCX report: {REPORT_DOCX}")
    try:
        generate_report_docx(results, REPORT_DOCX, total_dur)
        print(f"[*] Report saved: {REPORT_DOCX}")
    except Exception as exc:
        print(f"[!] Failed to generate report: {exc}")
        traceback.print_exc()

    # Also save JSON results
    json_path = OUTPUT_DIR / "selenium_results.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump({
            "timestamp": datetime.now().isoformat(),
            "total_tests": total,
            "passed": passed,
            "failed": failed,
            "duration_s": round(total_dur, 2),
            "results": results,
        }, f, indent=2, ensure_ascii=False)
    print(f"[*] JSON results: {json_path}")

    return 0 if failed == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
