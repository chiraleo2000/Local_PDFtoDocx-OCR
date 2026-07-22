# pylint: disable=broad-exception-caught
"""Golden regression: demo PDF vs Expected-output-testocr-demon.docx.

Runs with the exact UI settings the user specified:
  language=tha+eng, quality=fast, page=A4, margins=Moderate, trim=0, layout=flow

Skipped unless RUN_GOLDEN=1 (needs real models / GPU container).
"""
from __future__ import annotations

import os
import re
import zipfile
from difflib import SequenceMatcher
from pathlib import Path

import pytest

pytest.importorskip("docx")
from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

ROOT = Path(__file__).resolve().parents[1]
PDF = ROOT / "tests" / "fixtures" / "testocrtor-demo.pdf"
if not PDF.exists():
    PDF = ROOT / "tests" / "testocrtor-demo.pdf"
GOLD = ROOT / "tests" / "Expected-output-testocr-demon.docx"

_THAI_RE = re.compile(r"[\u0E00-\u0E7F]")
_KEEP_RE = re.compile(r"[\u0E00-\u0E7F0-9a-zA-Z@.]+")

REQUIRED_NEEDLES = [
    "แผนภูมิ",
    "ราชการ",
    "364",
    "931",
    "3.1",
    "2.1",
    "2.4",
]


def _norm(text: str) -> str:
    text = re.sub(r"(?:วรร[ทณ]|บรรณาธิการ|ฯลฯ)+", "", text or "")
    text = re.sub(r"\s+", "", text)
    return "".join(_KEEP_RE.findall(text)).lower()


def _docx_plain(path: Path) -> str:
    d = Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _docx_stats(path: Path) -> dict:
    d = Document(str(path))
    frames = d.element.body.findall(".//" + qn("w:framePr"))
    with zipfile.ZipFile(path) as z:
        images = [n for n in z.namelist() if n.startswith("word/media/")]
    text = _docx_plain(path)
    return {
        "tables": len(d.tables),
        "images": len(images),
        "frames": len(frames),
        "thai": len(_THAI_RE.findall(text)),
        "text": text,
    }


@pytest.mark.skipif(
    os.getenv("RUN_GOLDEN", "0").strip() not in ("1", "true", "yes"),
    reason="Set RUN_GOLDEN=1 to run model-backed golden demo",
)
@pytest.mark.skipif(not PDF.exists() or not GOLD.exists(),
                    reason="demo PDF or Expected DOCX missing")
def test_golden_demo_matches_expected(tmp_path):
    os.environ["LAYOUT_MODE"] = "flow"
    os.environ["DOCLING_REOCR"] = "1"
    os.environ["DOCLING_SPARSE_RECOVERY"] = "text"
    os.environ["TABLE_ENGINE"] = "opencv"
    os.environ["DISABLE_TROCR_PRELOAD"] = "0"

    from src.pipeline import OCRPipeline

    pipe = OCRPipeline()
    res = pipe.process_pdf(
        str(PDF),
        quality="fast",
        languages="tha+eng",
        header_trim=0,
        footer_trim=0,
        page_size="A4",
        margin_preset="Moderate",
        layout_mode="flow",
    )
    assert res["success"], res.get("error")
    out = Path(res["files"]["docx"])
    assert out.exists()

    gold = _docx_stats(GOLD)
    actual = _docx_stats(out)

    assert actual["frames"] == 0, (
        f"Expected flowing DOCX (framePr=0), got {actual['frames']}")
    assert actual["tables"] == gold["tables"] == 2, (
        f"tables actual={actual['tables']} gold={gold['tables']}")
    assert actual["images"] == gold["images"] == 2, (
        f"images actual={actual['images']} gold={gold['images']}")
    assert 1800 <= actual["thai"] <= 3200, (
        f"thai chars out of band: {actual['thai']}")

    gn, an = _norm(gold["text"]), _norm(actual["text"])
    sim = SequenceMatcher(None, gn, an).ratio()
    g_tok = set(re.findall(r"[\u0E00-\u0E7F]{2,}|\d{2,}", gold["text"]))
    a_tok = set(re.findall(r"[\u0E00-\u0E7F]{2,}|\d{2,}", actual["text"]))
    jacc = (len(g_tok & a_tok) / len(g_tok | a_tok)
            if (g_tok or a_tok) else 0.0)
    missing = [n for n in REQUIRED_NEEDLES if n not in actual["text"]]
    assert not missing, f"Missing required needles: {missing}"
    assert sim >= 0.65 or jacc >= 0.22, (
        f"Text similarity too low vs Expected "
        f"(sim={sim:.3f}, jaccard={jacc:.3f})")
