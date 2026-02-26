"""
Backend unit / integration tests for the OCR pipeline — v1.0.
Tests the pipeline, services, preprocessor, OCR engine, exporter,
CorrectionStore + manual-region APIs, and security hardening.
"""
import os
import shutil
import tempfile
from pathlib import Path

import numpy as np
import pytest

# ── Fixtures ──────────────────────────────────────────────────────────────────

TESTS_DIR = Path(__file__).parent
TEST_PDF = TESTS_DIR / "testocrtor.pdf"


@pytest.fixture(scope="module")
def pipeline():
    from src.pipeline import OCRPipeline
    return OCRPipeline()


@pytest.fixture(scope="module")
def pipeline_result(pipeline):
    assert TEST_PDF.exists(), f"Test PDF not found at {TEST_PDF}"
    result = pipeline.process_pdf(str(TEST_PDF), quality="fast")
    return result


# ══════════════════════════════════════════════════════════════════════════════
# Pipeline tests
# ══════════════════════════════════════════════════════════════════════════════

class TestOCRPipeline:
    def test_process_pdf_success(self, pipeline_result):
        """Pipeline should succeed on the test PDF."""
        assert pipeline_result["success"] is True
        assert pipeline_result["error"] is None

    def test_process_pdf_has_text(self, pipeline_result):
        """Extracted text should be non-empty."""
        assert isinstance(pipeline_result["text"], str)
        assert len(pipeline_result["text"]) > 0

    def test_process_pdf_metadata(self, pipeline_result):
        """Metadata should contain expected keys."""
        meta = pipeline_result["metadata"]
        assert "pages" in meta
        assert meta["pages"] >= 1
        assert "tables" in meta
        assert "figures" in meta

    def test_process_pdf_files_dict(self, pipeline_result):
        """Files dict should contain paths for txt and docx."""
        files = pipeline_result["files"]
        assert "txt" in files
        assert "docx" in files

    def test_process_pdf_txt_exists(self, pipeline_result):
        """Generated TXT file should exist on disk."""
        txt_path = pipeline_result["files"].get("txt")
        assert txt_path is not None
        assert os.path.exists(txt_path), f"TXT file not found: {txt_path}"
        assert os.path.getsize(txt_path) > 0

    def test_process_pdf_docx_exists(self, pipeline_result):
        """Generated DOCX file should exist on disk."""
        docx_path = pipeline_result["files"].get("docx")
        assert docx_path is not None
        assert os.path.exists(docx_path), f"DOCX file not found: {docx_path}"
        assert os.path.getsize(docx_path) > 0

    def test_process_pdf_docx_readable(self, pipeline_result):
        """Generated DOCX should be a valid Word document."""
        from docx import Document
        docx_path = pipeline_result["files"].get("docx")
        doc = Document(docx_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        assert len(paragraphs) > 0

    def test_process_pdf_invalid_path(self, pipeline):
        """Pipeline should return success=False for a missing file."""
        result = pipeline.process_pdf("/nonexistent/file.pdf")
        assert result["success"] is False
        assert result["error"] is not None

    def test_get_status(self, pipeline):
        """get_status should return a dict with expected keys."""
        status = pipeline.get_status()
        assert "ocr_engines" in status
        assert "easyocr" in status["ocr_engines"]

    def test_get_status_has_corrections(self, pipeline):
        """get_status should include correction stats (v0.5)."""
        status = pipeline.get_status()
        assert "corrections" in status
        c = status["corrections"]
        assert "total_manual_corrections" in c
        assert "next_retrain_at" in c
        assert "retrain_interval" in c


# ══════════════════════════════════════════════════════════════════════════════
# CorrectionStore tests (v0.5)
# ══════════════════════════════════════════════════════════════════════════════

class TestCorrectionStore:
    @pytest.fixture()
    def store(self, tmp_path):
        from src.correction_store import CorrectionStore
        return CorrectionStore(data_dir=str(tmp_path / "corr"), retrain_interval=5)

    def test_init_creates_dirs(self, store):
        assert store.images_dir.exists()
        assert store.labels_dir.exists()

    def test_log_correction_returns_dict(self, store):
        img = np.ones((200, 300, 3), dtype=np.uint8) * 128
        result = store.log_correction(
            page_image=img, bbox=[10, 20, 100, 80],
            region_class="table", page_number=0, pdf_name="test.pdf",
        )
        assert "correction_id" in result
        assert "total_corrections" in result
        assert result["total_corrections"] == 1

    def test_log_correction_saves_image(self, store):
        img = np.ones((200, 300, 3), dtype=np.uint8) * 128
        store.log_correction(img, [10, 20, 100, 80], "figure", 0, "t.pdf")
        pngs = list(store.images_dir.glob("*.png"))
        assert len(pngs) >= 1

    def test_log_correction_saves_label(self, store):
        img = np.ones((200, 300, 3), dtype=np.uint8) * 128
        store.log_correction(img, [10, 20, 100, 80], "table", 0, "t.pdf")
        txts = list(store.labels_dir.glob("*.txt"))
        assert len(txts) >= 1
        content = txts[0].read_text()
        # YOLO format: class cx cy w h
        parts = content.strip().split()
        assert len(parts) == 5
        assert parts[0] == "5"  # table class ID

    def test_corrections_jsonl_appended(self, store):
        img = np.ones((200, 300, 3), dtype=np.uint8) * 128
        store.log_correction(img, [10, 20, 100, 80], "figure", 0)
        store.log_correction(img, [30, 40, 150, 120], "table", 1)
        assert store.corrections_file.exists()
        lines = store.corrections_file.read_text().strip().split("\n")
        assert len(lines) == 2

    def test_get_stats(self, store):
        img = np.ones((200, 300, 3), dtype=np.uint8) * 128
        store.log_correction(img, [10, 20, 100, 80], "table", 0)
        stats = store.get_stats()
        assert stats["total_manual_corrections"] == 1
        assert stats["retrain_interval"] == 5
        assert stats["images_count"] >= 1

    def test_retrain_triggers_at_interval(self, store):
        """Retrain should trigger when correction count hits interval."""
        img = np.ones((200, 300, 3), dtype=np.uint8) * 128
        for i in range(4):
            r = store.log_correction(img, [10, 20, 100, 80], "table", i)
            assert r["retrain_triggered"] is False
        # 5th correction should trigger retrain (interval=5)
        r = store.log_correction(img, [10, 20, 100, 80], "table", 4)
        assert r["retrain_triggered"] is True

    def test_auto_source_not_counted(self, store):
        """Auto-detected regions should not count towards retrain trigger."""
        img = np.ones((200, 300, 3), dtype=np.uint8) * 128
        store.log_correction(img, [10, 20, 100, 80], "table", 0,
                             source="auto")
        assert store._correction_count == 0

    def test_get_corrections_log(self, store):
        img = np.ones((200, 300, 3), dtype=np.uint8) * 128
        store.log_correction(img, [10, 20, 100, 80], "table", 0, "a.pdf")
        log = store.get_corrections_log(limit=10)
        assert len(log) >= 1
        assert log[0]["class"] == "table"

    def test_log_page_detections(self, store):
        img = np.ones((200, 300, 3), dtype=np.uint8) * 128
        dets = {
            "tables": [{"bbox": [10, 20, 100, 80], "class": "table", "confidence": 0.9}],
            "figures": [{"bbox": [50, 60, 200, 180], "class": "figure", "confidence": 0.8}],
        }
        store.log_page_detections(img, dets, 0, "bulk.pdf")
        txts = [f for f in store.labels_dir.glob("*.txt") if "auto_" in f.name]
        assert len(txts) >= 1


# ══════════════════════════════════════════════════════════════════════════════
# Pipeline manual correction API tests (v0.5)
# ══════════════════════════════════════════════════════════════════════════════

class TestPipelineCorrections:
    def test_detect_page_regions(self, pipeline):
        """detect_page_regions should return detections for a valid PDF page."""
        assert TEST_PDF.exists()
        result = pipeline.detect_page_regions(str(TEST_PDF), 0)
        assert result["success"] is True
        assert "page_image" in result
        assert "detections" in result
        assert isinstance(result["detections"], dict)

    def test_detect_page_regions_invalid_page(self, pipeline):
        result = pipeline.detect_page_regions(str(TEST_PDF), 9999)
        assert result["success"] is False

    def test_add_manual_region(self, pipeline):
        """add_manual_region should log a correction and return stats."""
        img = np.ones((200, 300, 3), dtype=np.uint8) * 128
        result = pipeline.add_manual_region(
            page_image=img, bbox=[10, 20, 100, 80],
            region_class="table", page_number=0, pdf_name="test.pdf",
        )
        assert "correction_id" in result
        assert "total_corrections" in result

    def test_process_pdf_with_corrections(self, pipeline):
        """process_pdf_with_corrections should accept manual_regions."""
        assert TEST_PDF.exists()
        manual = {0: [{"bbox": [50, 50, 200, 200], "class": "table"}]}
        result = pipeline.process_pdf_with_corrections(
            str(TEST_PDF), manual_regions=manual, quality="fast",
        )
        assert result["success"] is True
        assert result["metadata"]["manual_corrections"] == 1

    def test_process_pdf_with_corrections_no_manual(self, pipeline):
        """Should work fine with no manual regions."""
        result = pipeline.process_pdf_with_corrections(
            str(TEST_PDF), quality="fast")
        assert result["success"] is True
        assert result["metadata"]["manual_corrections"] == 0


# ══════════════════════════════════════════════════════════════════════════════
# OCR Engine tests
# ══════════════════════════════════════════════════════════════════════════════

class TestOCREngine:
    def test_available_engines(self):
        from src.ocr_engine import OCREngine
        engine = OCREngine()
        engines = engine.get_available_engines()
        assert isinstance(engines, dict)
        assert "easyocr" in engines
        assert "paddleocr" in engines

    def test_easyocr_available(self):
        from src.ocr_engine import EASYOCR_AVAILABLE
        assert EASYOCR_AVAILABLE, "EasyOCR should be available in the test environment"

    def test_ocr_image(self):
        """OCR should return text from a simple image."""
        import numpy as np
        import cv2
        from src.ocr_engine import OCREngine

        engine = OCREngine()
        # Create a white image with black text pattern
        img = np.ones((100, 400), dtype=np.uint8) * 255
        cv2.putText(img, "Hello OCR", (10, 60), cv2.FONT_HERSHEY_SIMPLEX,
                    1.5, (0, 0, 0), 2)
        result = engine.ocr_image(img)
        assert isinstance(result, dict)
        assert "text" in result
        assert "confidence" in result
        assert "engine_used" in result

    def test_is_available(self):
        from src.ocr_engine import OCREngine
        engine = OCREngine()
        assert engine.is_available() is True


# ══════════════════════════════════════════════════════════════════════════════
# Preprocessor tests
# ══════════════════════════════════════════════════════════════════════════════

class TestOpenCVPreprocessor:
    def test_preprocess_fast(self):
        import numpy as np
        from src.preprocessor import OpenCVPreprocessor

        pp = OpenCVPreprocessor(quality="fast")
        img = np.ones((200, 300, 3), dtype=np.uint8) * 200
        out = pp.preprocess(img)
        assert out is not None
        assert out.size > 0

    def test_preprocess_balanced(self):
        import numpy as np
        from src.preprocessor import OpenCVPreprocessor

        pp = OpenCVPreprocessor(quality="balanced")
        img = np.ones((200, 300, 3), dtype=np.uint8) * 200
        out = pp.preprocess(img)
        assert out is not None
        assert out.ndim in (2, 3)

    def test_preprocess_accurate(self):
        import numpy as np
        from src.preprocessor import OpenCVPreprocessor

        pp = OpenCVPreprocessor(quality="accurate")
        img = np.ones((200, 300, 3), dtype=np.uint8) * 200
        out = pp.preprocess(img)
        assert out is not None


# ══════════════════════════════════════════════════════════════════════════════
# Exporter tests
# ══════════════════════════════════════════════════════════════════════════════

class TestDocumentExporter:
    SAMPLE_TEXT = "Hello, World!\n\nThis is a test document.\n## Section\nSome content."

    def test_create_txt(self):
        from src.exporter import DocumentExporter
        exp = DocumentExporter()
        path = exp.create_txt(self.SAMPLE_TEXT, metadata="Test metadata")
        assert os.path.exists(path)
        content = Path(path).read_text(encoding="utf-8")
        assert "Hello, World!" in content
        assert "Test metadata" in content
        os.unlink(path)

    def test_create_docx(self):
        from src.exporter import DocumentExporter
        from docx import Document
        exp = DocumentExporter()
        path = exp.create_docx(self.SAMPLE_TEXT)
        assert path is not None
        assert os.path.exists(path)
        doc = Document(path)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Hello, World!" in full_text
        os.unlink(path)

    def test_create_html(self):
        from src.exporter import DocumentExporter
        exp = DocumentExporter()
        path = exp.create_html(self.SAMPLE_TEXT)
        assert os.path.exists(path)
        content = Path(path).read_text(encoding="utf-8")
        assert "<!DOCTYPE html>" in content
        assert "Hello, World!" in content
        os.unlink(path)

    def test_create_all(self):
        from src.exporter import DocumentExporter
        exp = DocumentExporter()
        files = exp.create_all(self.SAMPLE_TEXT)
        assert "txt" in files
        assert "docx" in files
        assert "html" in files
        for fmt, path in files.items():
            if path:
                assert os.path.exists(path), f"{fmt} file not found"
                os.unlink(path)


# ══════════════════════════════════════════════════════════════════════════════
# Auth & History tests
# ══════════════════════════════════════════════════════════════════════════════

class TestAuthManager:
    def test_default_guest_login(self):
        from src.services import AuthManager
        auth = AuthManager()
        result = auth.login("guest", "guest123")
        assert result["success"] is True
        assert "token" in result
        assert result["username"] == "guest"

    def test_invalid_login(self):
        from src.services import AuthManager
        auth = AuthManager()
        result = auth.login("guest", "wrongpassword")
        assert result["success"] is False

    def test_register_and_login(self):
        import uuid
        from src.services import AuthManager
        auth = AuthManager()
        uname = f"testuser_{uuid.uuid4().hex[:8]}"
        reg = auth.register(uname, "pass1234")
        assert reg["success"] is True
        login = auth.login(uname, "pass1234")
        assert login["success"] is True

    def test_session_validate(self):
        from src.services import AuthManager
        auth = AuthManager()
        login = auth.login("guest", "guest123")
        token = login["token"]
        username = auth.validate_session(token)
        assert username == "guest"

    def test_logout_invalidates_session(self):
        from src.services import AuthManager
        auth = AuthManager()
        login = auth.login("guest", "guest123")
        token = login["token"]
        auth.logout(token)
        assert auth.validate_session(token) is None


class TestHistoryManager:
    def test_save_and_list(self, tmp_path, monkeypatch):
        import tempfile
        from src.services import HistoryManager

        monkeypatch.setenv("HISTORY_RETENTION_DAYS", "30")
        # Patch tempdir so history goes into tmp_path
        monkeypatch.setattr(tempfile, "gettempdir", lambda: str(tmp_path))

        hist = HistoryManager()
        # Create dummy output files
        txt_file = tmp_path / "out.txt"
        docx_file = tmp_path / "out.docx"
        txt_file.write_text("hello")
        docx_file.write_bytes(b"PK\x03\x04")  # minimal docx/zip header

        entry_id = hist.save_result(
            "testuser", "sample.pdf",
            {"txt": str(txt_file), "docx": str(docx_file)},
            {"pages": 1}
        )
        assert entry_id is not None

        entries = hist.list_entries("testuser")
        assert len(entries) >= 1
        assert entries[0]["original_filename"] == "sample.pdf"

    def test_get_file_path(self, tmp_path, monkeypatch):
        import tempfile
        from src.services import HistoryManager

        monkeypatch.setattr(tempfile, "gettempdir", lambda: str(tmp_path))

        hist = HistoryManager()
        txt_file = tmp_path / "out2.txt"
        docx_file = tmp_path / "out2.docx"
        txt_file.write_text("text content")
        docx_file.write_bytes(b"PK\x03\x04")

        entry_id = hist.save_result(
            "user2", "doc.pdf",
            {"txt": str(txt_file), "docx": str(docx_file)},
        )
        docx_path = hist.get_file_path("user2", entry_id, "docx")
        txt_path = hist.get_file_path("user2", entry_id, "txt")
        assert docx_path is not None and os.path.exists(docx_path)
        assert txt_path is not None and os.path.exists(txt_path)


# ══════════════════════════════════════════════════════════════════════════════
# Security tests (v1.0)
# ══════════════════════════════════════════════════════════════════════════════

class TestSecurityHardening:
    """Verify security measures introduced in v1.0."""

    def test_pipeline_rejects_non_pdf(self, pipeline):
        """Pipeline should reject files without .pdf extension."""
        result = pipeline.process_pdf("document.txt")
        assert result["success"] is False
        assert ".pdf" in result["error"].lower()

    def test_pipeline_rejects_empty_path(self, pipeline):
        """Pipeline should reject empty/None paths."""
        result = pipeline.process_pdf("")
        assert result["success"] is False

    def test_pipeline_validate_trim_bounds(self, pipeline):
        """Header/footer trim should be clamped to [0, 25]."""
        # This shouldn't crash — values get clamped internally
        result = pipeline.process_pdf(
            str(TEST_PDF), quality="fast",
            header_trim=999, footer_trim=-5)
        # Should still succeed (values are clamped)
        assert result["success"] is True

    def test_auth_rejects_short_password(self):
        """Registration should reject passwords shorter than 6 chars."""
        from src.services import AuthManager
        auth = AuthManager()
        result = auth.register("validuser", "ab")
        assert result["success"] is False

    def test_auth_rejects_invalid_username(self):
        """Registration should reject usernames with special chars."""
        from src.services import AuthManager
        auth = AuthManager()
        result = auth.register("../evil", "password123")
        assert result["success"] is False

    def test_correction_store_sanitizes_filename(self):
        """CorrectionStore should sanitize pdf_name in filenames."""
        from src.correction_store import CorrectionStore
        store = CorrectionStore(data_dir=str(Path(tempfile.mkdtemp()) / "sec"))
        img = np.ones((200, 300, 3), dtype=np.uint8) * 128
        result = store.log_correction(
            page_image=img, bbox=[10, 20, 100, 80],
            region_class="table", page_number=0,
            pdf_name="../../etc/passwd.pdf",
        )
        # The correction_id should not contain path traversal
        assert ".." not in result["correction_id"]
        assert "/" not in result["correction_id"]

    def test_detect_page_regions_validates_path(self, pipeline):
        """detect_page_regions should validate PDF path."""
        result = pipeline.detect_page_regions("/nonexistent.pdf", 0)
        assert result["success"] is False

    def test_process_with_corrections_validates_path(self, pipeline):
        """process_pdf_with_corrections should validate PDF path."""
        result = pipeline.process_pdf_with_corrections("/nonexistent.pdf")
        assert result["success"] is False
