"""Regression tests for Docling adapter + table-drop fixes."""
from types import SimpleNamespace
from unittest.mock import MagicMock

import numpy as np
import pytest

from src.docling_adapter import (
    docling_to_blocks,
    detections_from_docling,
    _suppress_text_in_structure,
    _looks_garbled_for_thai,
    _has_usable_thai,
    _ocr_text_block,
    _thai_reocr_enabled,
)
from src.pipeline import ContentBlock
from src.exporter import DocumentExporter

docx = pytest.importorskip("docx")
from docx import Document  # noqa: E402


class _BBox:
    def __init__(self, l, t, r, b):
        self.l, self.t, self.r, self.b = l, t, r, b

    def to_top_left_origin(self, page_h):
        return self


class _Prov:
    def __init__(self, page_no, bbox):
        self.page_no = page_no
        self.bbox = bbox


class _Label:
    def __init__(self, value):
        self.value = value


def _fake_docling_doc():
    """Minimal DoclingDocument-like object with text/table/picture."""
    text = SimpleNamespace(
        label=_Label("text"),
        text="วิเคราะห์มาตรฐานการผลิตหม่อนไหม",
        prov=[_Prov(1, _BBox(50, 40, 400, 80))],
    )
    table = SimpleNamespace(
        label=_Label("table"),
        text="",
        prov=[_Prov(1, _BBox(50, 120, 500, 300))],
        data=SimpleNamespace(
            num_cols=2,
            grid=[
                [SimpleNamespace(text="ชื่อ"), SimpleNamespace(text="จำนวน")],
                [SimpleNamespace(text="ข้าว"), SimpleNamespace(text="10")],
            ],
        ),
        export_to_html=lambda: (
            "<table><tr><th>ชื่อ</th><th>จำนวน</th></tr>"
            "<tr><td>ข้าว</td><td>10</td></tr></table>"
        ),
    )
    picture = SimpleNamespace(
        label=_Label("picture"),
        text="",
        prov=[_Prov(1, _BBox(60, 340, 260, 480))],
        image=None,
    )
    page = SimpleNamespace(size=SimpleNamespace(width=600, height=800))
    doc = SimpleNamespace(
        pages={1: page},
        texts=[text],
        tables=[table],
        pictures=[picture],
    )
    doc.iterate_items = lambda: [
        (text, 0), (table, 0), (picture, 0)]
    return doc


def test_docling_adapter_maps_table_figure_text():
    doc = _fake_docling_doc()
    blocks = docling_to_blocks(doc, ocr=None, page_images={}, languages="tha+eng")
    types = {b.block_type for b in blocks}
    assert "table" in types
    assert "figure" in types
    assert "text" in types
    table = next(b for b in blocks if b.block_type == "table")
    assert "<table>" in table.table_html
    assert "ข้าว" in table.table_html
    assert table.bbox == [50, 120, 500, 300]
    n_tables = sum(1 for b in blocks if b.block_type == "table")
    n_figures = sum(1 for b in blocks if b.block_type == "figure")
    assert n_tables == 1
    assert n_figures == 1


def test_detections_from_docling_match_adapter_counts():
    doc = _fake_docling_doc()
    dets = detections_from_docling(doc, page_idx=0)
    assert len(dets["tables"]) == 1
    assert len(dets["figures"]) == 1
    assert len(dets["text_regions"]) >= 1


def test_suppress_text_inside_table():
    blocks = [
        ContentBlock("text", 0, 130, 60, text="cell junk",
                     bbox=[60, 130, 200, 160], page_width=600, page_height=800),
        ContentBlock("table", 0, 120, 50, text="A\tB",
                     table_html="<table><tr><td>A</td><td>B</td></tr></table>",
                     bbox=[50, 120, 500, 300], page_width=600, page_height=800),
        ContentBlock("text", 0, 40, 50, text="Title outside",
                     bbox=[50, 40, 400, 80], page_width=600, page_height=800),
    ]
    out = _suppress_text_in_structure(blocks)
    texts = [b.text for b in out if b.block_type == "text"]
    assert "Title outside" in texts
    assert "cell junk" not in texts


def test_suppress_text_inside_figure():
    """Org-chart labels overlapping a figure must be dropped (gold = image only)."""
    blocks = [
        ContentBlock(
            "figure", 0, 200, 50,
            figure={"base64": "abc", "width": 200, "height": 200},
            bbox=[50, 200, 500, 600], page_width=600, page_height=800),
        ContentBlock(
            "text", 0, 250, 80, text="กรมหม่อนไหม org junk",
            bbox=[80, 250, 400, 500], page_width=600, page_height=800),
        ContentBlock(
            "text", 0, 40, 50, text="7) แผนภูมิการแบ่งส่วนราชการ",
            bbox=[50, 40, 400, 80], page_width=600, page_height=800),
    ]
    out = _suppress_text_in_structure(blocks)
    texts = [b.text for b in out if b.block_type == "text"]
    assert "7) แผนภูมิการแบ่งส่วนราชการ" in texts
    assert not any("กรมหม่อนไหม" in t for t in texts)


def test_polish_inventory_cleans_crumbs_and_headers():
    from src.docling_adapter import _polish_inventory_table
    html = (
        "<table>"
        "<tr><th>รายการ</th><th>รายละเอยด</th><th></th></tr>"
        "<tr><td>เครื่องคอมพิวเตอร์แม่ข่าย</td><td>แบบที่ 2</td><td>16</td></tr>"
        "<tr><td>84</td><td>- เครื่องคอมพิวเตอร์โน้ตบุ๊ก</td><td>166</td></tr>"
        "<tr><td>ชุดโปรแกรมระบบปฏิบัติการ</td><td>Microsoft Windows</td><td>165</td></tr>"
        "</table>"
    )
    out_html, out_text = _polish_inventory_table(html, "")
    assert "จำนวน" in out_html
    assert "รายละเอียด" in out_html
    assert ">84<" not in out_html
    assert "ฮาร์ดแวร์" in out_html
    assert "ซอฟต์แวร์" in out_html


def test_plausible_keeps_latin_product_names():
    from src.docling_adapter import (
        _plausible_table_cell, _looks_like_healthy_latin, _prefer_digit_seed,
    )
    assert _looks_like_healthy_latin("Microsoft Windows")
    assert _looks_like_healthy_latin("ESET")
    assert _plausible_table_cell("Microsoft Windows")
    assert _plausible_table_cell("Next Generation Firewall")
    assert _prefer_digit_seed("Microsoft Windows", "คล่องแคล่ว") == "Microsoft Windows"


def test_polish_staff_table_fixes_header_crumb():
    from src.docling_adapter import _polish_staff_table
    html = (
        "<table>"
        "<tr><th>แขน</th><th>จำนวน(คน)</th></tr>"
        "<tr><td>ข้าราชการ</td><td>364</td></tr>"
        "<tr><td>0</td><td>349</td></tr>"
        "<tr><td>รวมทั้งสิ้น</td><td>931</td></tr>"
        "</table>"
    )
    out_html, _ = _polish_staff_table(html, "")
    assert "ตำแหน่ง" in out_html
    assert ">0<" not in out_html


def test_pick_better_table_prefers_taller_grid():
    from src.docling_adapter import _pick_better_table
    # Collapsed 13-row Docling vs taller OpenCV-like plain
    html_a = "<table>" + "".join(
        f"<tr><td>x{i}</td><td>y{i}</td><td>{i}</td></tr>"
        for i in range(13)) + "</table>"
    # Sparse Thai in A
    text_a = "\n".join(f"x{i}\ty{i}\t{i}" for i in range(13))
    rows_b = [
        "รายการ\tรายละเอียด\tจำนวน",
        "ฮาร์ดแวร์\t\t",
        "เครื่องคอมพิวเตอร์แม่ข่าย\tแบบที่ 2\t16",
    ] + [f"รายการ{i}\tรายละเอียด{i}\t{i}" for i in range(20)]
    text_b = "\n".join(rows_b)
    html_b = "<table>" + "".join(
        "<tr>" + "".join(f"<td>{c}</td>" for c in r.split("\t")) + "</tr>"
        for r in rows_b) + "</table>"
    html, text = _pick_better_table(html_a, text_a, html_b, text_b)
    assert "ฮาร์ดแวร์" in (html + text)
    assert html.count("<tr") >= 20


def test_table_quality_weak_collapsed_inventory():
    from src.docling_adapter import _table_quality_weak
    assert _table_quality_weak({
        "thai": 150, "left_thai_fill": 0.40, "fill": 0.50,
        "nrows": 13, "ncols": 3, "cells": 39, "filled": 20,
    })
    assert not _table_quality_weak({
        "thai": 400, "left_thai_fill": 0.70, "fill": 0.80,
        "nrows": 25, "ncols": 3, "cells": 75, "filled": 60,
    })


def test_looks_like_thai_hallucination_rejects_easyocr_junk():
    from src.docling_adapter import _looks_like_thai_hallucination
    assert _looks_like_thai_hallucination("ยอมแพ้")
    assert _looks_like_thai_hallucination("สังคมพิสูจน์มักฯ '_")
    assert not _looks_like_thai_hallucination("เครื่องคอมพิวเตอร์แม่ข่าย")


def test_exporter_rebuilds_table_from_plain_text():
    """Absolute export must not drop tables that lack table_html."""
    blocks = [
        ContentBlock(
            "table", 0, 100, 50,
            text="Name\tQty\nRice\t10",
            table_html="",
            bbox=[50, 100, 500, 300],
            page_width=600, page_height=800,
        ),
    ]
    exp = DocumentExporter()
    files = exp.create_all_from_blocks(blocks, "Pages: 1", "A4", "Normal",
                                       layout_mode="absolute")
    doc = Document(files["docx"])
    assert len(doc.tables) == 1
    flat = "\n".join(c.text for row in doc.tables[0].rows for c in row.cells)
    assert "Rice" in flat or "Name" in flat


def test_layout_backend_env_default(monkeypatch):
    monkeypatch.delenv("LAYOUT_BACKEND", raising=False)
    from src.docling_backend import layout_backend
    assert layout_backend() == "docling"
    monkeypatch.setenv("LAYOUT_BACKEND", "yolo")
    assert layout_backend() == "yolo"


def test_garbled_latin_detected():
    # RapidOCR-on-Thai soup must be rejected
    assert _looks_garbled_for_thai(
        "COMMSSUBLMACLUNGMUNEUSLUOBLUMLABEMUI ENUCSH")
    # Real English / product names must be KEPT on tha+eng jobs
    assert not _looks_garbled_for_thai("Hello English only")
    assert not _looks_garbled_for_thai("Microsoft Windows Server")
    assert not _looks_garbled_for_thai(
        "วิเคราะห์และตรวจสอบมาตรฐานการผลิต")
    assert _has_usable_thai("ชื่อ\tจำนวน\nข้าว\t10", min_chars=3)
    assert not _has_usable_thai("A B C ข้าว", min_chars=6, min_density=0.25)


def test_thai_reocr_force_off(monkeypatch):
    monkeypatch.setenv("DOCLING_REOCR", "force-off")
    assert _thai_reocr_enabled("tha+eng") is False
    monkeypatch.setenv("DOCLING_REOCR", "0")
    assert _thai_reocr_enabled("tha+eng") is True


def test_ocr_text_block_refuses_latin_fallback():
    """Failed Thai-TrOCR must not keep RapidOCR Latin garbage."""
    ocr = MagicMock()
    ocr.ocr_image.return_value = {"text": "", "lines": []}
    ocr.ocr_full_page.return_value = {"text": "", "lines": []}
    img = np.zeros((200, 400, 3), dtype=np.uint8)
    block = _ocr_text_block(
        ocr, img, [10, 10, 200, 40], 0, 400, 200, "tha+eng",
        fallback_text="COMMSSUBLMACLUNGMUNEUSLUOBLUMLABEMUI")
    assert block.text.strip() == ""
    assert not block.lines


def test_section_marker_rejects_chart_decimals():
    from src.pipeline import _is_valid_section_marker
    assert _is_valid_section_marker("2.1")
    assert _is_valid_section_marker("11)")
    assert not _is_valid_section_marker("0.78")
    assert not _is_valid_section_marker("0.5")
    assert not _is_valid_section_marker("99.99")


def test_dedup_blocks_collapses_near_and_marker_dups():
    from src.pipeline import ContentBlock, _dedup_blocks
    a = ContentBlock(
        block_type="text", page=0, y_top=10, x_left=0,
        text="4) จัดทำมาตรฐานการผลิตสินค้าหม่อนไหม",
        bbox=[0, 10, 400, 40], page_width=500, page_height=700)
    b = ContentBlock(
        block_type="text", page=0, y_top=12, x_left=0,
        text="4) จัดทำมาตรฐานการผลิตสินค้าหม่อนไหมและผลิตภัณฑ์",
        bbox=[0, 12, 420, 42], page_width=500, page_height=700)
    c = ContentBlock(
        block_type="text", page=0, y_top=200, x_left=0,
        text="7) แผนภูมิการแบ่งส่วนราชการ",
        bbox=[0, 200, 400, 230], page_width=500, page_height=700)
    d = ContentBlock(
        block_type="text", page=0, y_top=80, x_left=0,
        text="7) ดำเนินการเกี่ยวกับฐานข้อมูล",
        bbox=[0, 80, 400, 110], page_width=500, page_height=700)
    out = _dedup_blocks([a, b, d, c])
    texts = [x.text for x in out]
    assert len([t for t in texts if t.startswith("4)")]) == 1
    assert any("ผลิตภัณฑ์" in t for t in texts)  # richer 4) kept
    assert any("แผนภูมิ" in t for t in texts)
    assert any("ฐานข้อมูล" in t for t in texts)
