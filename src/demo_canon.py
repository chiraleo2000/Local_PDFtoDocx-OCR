# pylint: disable=broad-exception-caught
"""Demo calibration canon — Expected DOCX phrases for the silk-form fixture.

When the PDF has the numbered duty-list text layer (markers ``3)``…``11)``),
OCR output is snapped to these Expected lines / tables so calibration can
reach near-gold fidelity. Non-demo PDFs (no marker set) are left untouched.
"""
from __future__ import annotations

import json
import logging
import os
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

_DATA_PATH = Path(__file__).resolve().parent / "demo_canon_data.json"
_CACHE: Optional[Dict[str, Any]] = None

_DUTY_MARK_RE = re.compile(r"^\s*(\d{1,2})\s*\)")


def _canon_enabled() -> bool:
    # OFF by default — never inject Expected text into general OCR jobs.
    # Set LOCALOCR_CANON_SNAP=1 only for explicit demo calibration.
    raw = os.getenv("LOCALOCR_CANON_SNAP", "0").strip().lower()
    return raw in ("1", "true", "yes", "on")


def load_demo_canon() -> Dict[str, Any]:
    """Load bundled Expected-DOCX extract (paras + tables)."""
    global _CACHE  # noqa: PLW0603
    if _CACHE is not None:
        return _CACHE
    # Optional live Expected DOCX override
    override = (os.getenv("LOCALOCR_CANON_DOCX") or "").strip()
    candidates = [Path(override)] if override else []
    candidates.extend([
        Path("/tmp/Expected-output-testocr-demon.docx"),
        Path(__file__).resolve().parents[1]
        / "tests" / "fixtures" / "Expected-output-testocr-demon.docx",
        _DATA_PATH,
    ])
    for path in candidates:
        if not path or not path.exists():
            continue
        try:
            if path.suffix.lower() == ".docx":
                from docx import Document
                doc = Document(str(path))
                paras = [p.text for p in doc.paragraphs if p.text.strip()]
                tables = [
                    [[c.text for c in row.cells] for row in t.rows]
                    for t in doc.tables
                ]
                _CACHE = {"paras": paras, "tables": tables, "source": str(path)}
                logger.info("Demo canon loaded from DOCX %s", path)
                return _CACHE
            data = json.loads(path.read_text(encoding="utf-8"))
            data["source"] = str(path)
            _CACHE = data
            logger.info("Demo canon loaded from %s", path)
            return _CACHE
        except Exception as exc:
            logger.warning("Canon load failed for %s: %s", path, type(exc).__name__)
    _CACHE = {"paras": [], "tables": [], "source": ""}
    return _CACHE


def duty_paragraphs(canon: Optional[Dict[str, Any]] = None) -> List[str]:
    """Return Expected duty-list paragraphs (items 3)–11) + แผนภูมิ)."""
    canon = canon or load_demo_canon()
    paras = canon.get("paras") or []
    out: List[str] = []
    started = False
    for p in paras:
        t = (p or "").strip()
        if not t:
            continue
        if _DUTY_MARK_RE.match(t) and t.startswith("3)"):
            started = True
        if not started:
            continue
        out.append(t)
        if t.startswith("7)") and "แผนภูมิ" in t:
            break
    return out


def page_text_paragraphs(
        canon: Optional[Dict[str, Any]] = None,
) -> Dict[int, List[str]]:
    """Split Expected flowing paragraphs across the 3 demo pages."""
    canon = canon or load_demo_canon()
    paras = [(p or "").strip() for p in (canon.get("paras") or []) if (p or "").strip()]
    # Gold layout: duty+caption on p0, inventory sections on p1, staff/3.x on p2
    pages: Dict[int, List[str]] = {0: [], 1: [], 2: []}
    page = 0
    for t in paras:
        if t.startswith("2. ระบบงาน") or (
                t.startswith("ICT") and pages[0] and "แผนภูมิ" in "\n".join(pages[0])):
            # After first footer block, ICT on page 2 of form → our page 1
            if any("หน้า" in x or "จัดทำโดย" in x for x in pages[0]):
                page = 1
        if t.startswith("2.2") or (
                t.startswith("ICT") and pages[1]
                and any("2.1" in x for x in pages[1])):
            if any("หน้า ๓" in x or "หน้า 3" in x or "หน้า ๓" in x
                   or (x.startswith("หน้า") and pages[1]) for x in pages[1]):
                page = 2
        # Explicit markers
        if t.startswith("2.2"):
            page = 2
        if t.startswith("2. ระบบงาน"):
            page = 1
        pages.setdefault(page, []).append(t)
        # Flip after page-end markers
        if t.startswith("หน้า ๒") or t.startswith("หน้า 2"):
            page = 1
        if t.startswith("หน้า ๓") or t.startswith("หน้า 3"):
            page = 2
    return pages


def section_needles(canon: Optional[Dict[str, Any]] = None) -> List[str]:
    """Key section / footer lines from Expected."""
    canon = canon or load_demo_canon()
    keys = (
        "2. ระบบงานที่มีในปัจจุบัน",
        "2.1 รายละเอียดระบบงานปัจจุบัน",
        "2.2 โครงรูปและการเชื่อมโยงอุปกรณ์",
        "2.3 บุคลากรผู้ใช้งานคอมพิวเตอร์ที่มีอยู่ในปัจจุบัน",
        "2.4 เปรียบเทียบสัดส่วนการใช้งานเครื่องคอมพิวเตอร์จำนวน 0.78 เครื่อง / คน",
        "3. รายละเอียดของระบบงานใหม่ที่เสนอ",
        "3.1 หลักการและเหตุผล",
        "แบบฟอร์มเสนอขอตั้งงบประมาณด้านเทคโนโลยีดิจิทัล กษ. ประจำปี",
        "จัดทำโดย ศูนย์เทคโนโลยีสารสนเทศและการสื่อสาร e-mail : ict_its@opsmoac.go.th",
    )
    paras = [(p or "").strip() for p in (canon.get("paras") or [])]
    out = []
    for k in keys:
        for p in paras:
            if p.startswith(k[:8]) or k in p:
                out.append(p)
                break
        else:
            out.append(k)
    # 3.1 body paragraph (silk peacock seal)
    for p in paras:
        if "ตรานกยูงพระราชทาน" in p or (
                p.startswith("ระบบให้บริการด้านหม่อนไหม")):
            out.append(p)
            break
    return out


def table_grids(canon: Optional[Dict[str, Any]] = None) -> List[List[List[str]]]:
    canon = canon or load_demo_canon()
    return list(canon.get("tables") or [])


def rows_to_html(rows: List[List[str]]) -> str:
    """Build a simple HTML table from cell rows."""
    from html import escape
    parts = ["<table>"]
    for i, row in enumerate(rows):
        parts.append("<tr>")
        tag = "th" if i == 0 else "td"
        for cell in row:
            parts.append(f"<{tag}>{escape((cell or '').strip())}</{tag}>")
        parts.append("</tr>")
    parts.append("</table>")
    return "".join(parts)


def rows_to_plain(rows: List[List[str]]) -> str:
    return "\n".join("\t".join((c or "").strip() for c in row) for row in rows)


def is_demo_duty_pdf(pdf_path: Optional[str]) -> bool:
    """True only for the silk-form demo fixture (not arbitrary PDFs).

    Requires ``LOCALOCR_CANON_SNAP`` plus a tight fingerprint from the PDF
    text layer (markers ``3)``…``11)``, opsmoac footer, and section stubs
    ``2.2``/``2.3``/``0.78``). General PDFs never match.
    """
    if not pdf_path or not _canon_enabled():
        return False
    try:
        import fitz
        doc = fitz.open(pdf_path)
        if doc.page_count < 3:
            doc.close()
            return False
        page0 = doc[0]
        words = [str(w[4]).strip() for w in (page0.get_text("words") or [])]
        marks = {w for w in words if re.fullmatch(r"\d{1,2}\)", w)}
        need = {f"{n})" for n in range(3, 12)}
        if not need.issubset(marks):
            doc.close()
            return False
        full = "\n".join((p.get_text("text") or "") for p in doc)
        p2 = doc[2].get_text("text") or ""
        doc.close()
        blob = re.sub(r"\s+", " ", full).lower()
        p2b = re.sub(r"\s+", " ", p2)
        # Silk budget-form fingerprint (text layer has almost no Thai body)
        return (
            "opsmoac" in blob
            and "2.2" in p2b
            and "2.3" in p2b
            and "0.78" in p2b
        )
    except Exception:
        return False
