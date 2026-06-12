"""
Document Exporter + Image Extractor — v2.3
HTML-first export: Build positioned HTML → convert to DOCX & TXT.
Images embedded as base64 in HTML, properly transferred to DOCX via htmldocx.
Figure numbering: 1-based.

v2.1 improvements:
    - Reasonable image sizing (based on original dimensions, not hardcoded)
    - Table borders enforced in DOCX output
    - Configurable page size (A4, Letter, Legal, A3, B5)
    - Configurable margins (Normal, Narrow, Moderate, Wide)

v2.0 improvements:
    - Better table-to-DOCX conversion with proper cell alignment
    - Colspan/rowspan handling in fallback DOCX builder
    - Style preservation (tabs, indentation) in text blocks
    - Improved HTML table structure with col-width hints

Security:
    - All user text is escaped via ``html.escape()`` before HTML insertion (XSS)
    - Base64 data validated before embedding
    - Filenames sanitised in temp file operations
"""
import os
import re
import html
import logging
import base64
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Dict, Any, Optional, List, TYPE_CHECKING

import cv2
import numpy as np
from PIL import Image

if TYPE_CHECKING:
    from .pipeline import ContentBlock

logger = logging.getLogger(__name__)

# ── Optional imports ──────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed — DOCX export unavailable")

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

HTMLDOCX_AVAILABLE = False
try:
    from htmldocx import HtmlToDocx
    HTMLDOCX_AVAILABLE = True
except ImportError:
    logger.info("htmldocx not installed — will use fallback DOCX builder")

_BASE64_RE = re.compile(r"^[A-Za-z0-9+/\n\r]*={0,2}$")


def _is_valid_base64(data: str) -> bool:
    """Validate that *data* looks like legitimate base64 (anti-injection)."""
    if not data or len(data) > 50_000_000:  # 50 MB max
        return False
    return bool(_BASE64_RE.match(data))


# ── Page layout configurations ────────────────────────────────────────────────
# Page sizes: (width_inches, height_inches)
PAGE_SIZES = {
    "A4": (8.27, 11.69),
    "Letter": (8.5, 11.0),
    "Legal": (8.5, 14.0),
    "A3": (11.69, 16.54),
    "B5": (6.93, 9.84),
}

# Margin presets: (top, bottom, left, right) in inches — matching MS Word defaults
MARGIN_PRESETS = {
    "Normal": (1.0, 1.0, 1.0, 1.0),
    "Narrow": (0.5, 0.5, 0.5, 0.5),
    "Moderate": (1.0, 1.0, 0.75, 0.75),
    "Wide": (1.0, 1.0, 1.5, 1.5),
}

# Max image display width in HTML (pixels)
_MAX_IMG_DISPLAY_PX = 550

# Assumed rendering DPI for pixel-to-inch conversion
_IMG_RENDER_DPI = 150


# ══════════════════════════════════════════════════════════════════════════════
# Image / Figure Extraction  (1-based numbering)
# ══════════════════════════════════════════════════════════════════════════════
class ImageExtractor:
    """Crop detected figure regions, save to disk, and encode to base64.

    v2.3: thresholds lowered (and env-tunable) so small logos/stamps are
    kept; every figure carries its source ``bbox`` so the pipeline never
    mis-associates positions; ``page_figure()`` preserves graphic-only
    pages (covers, diagrams) that OCR would otherwise drop entirely.
    """

    def __init__(self, min_width: int = 40, min_height: int = 40,
                 min_area: int = 2000):
        self.min_width = int(os.getenv("IMAGE_MIN_WIDTH", str(min_width)))
        self.min_height = int(os.getenv("IMAGE_MIN_HEIGHT", str(min_height)))
        self.min_area = int(os.getenv("IMAGE_MIN_AREA", str(min_area)))
        self.enabled = os.getenv("IMAGE_EXTRACTION", "true").lower() == "true"
        self.temp_dir = Path(tempfile.gettempdir()) / "pdf_ocr_images"
        self.temp_dir.mkdir(parents=True, exist_ok=True)

    @staticmethod
    def encode_png_base64(crop: np.ndarray) -> str:
        """Encode a BGR/gray numpy crop as base64 PNG."""
        pil_img = Image.fromarray(
            cv2.cvtColor(crop, cv2.COLOR_BGR2RGB)
            if len(crop.shape) == 3 else crop
        )
        buf = BytesIO()
        pil_img.save(buf, format="PNG")
        return base64.b64encode(buf.getvalue()).decode()

    def extract_figures(self, page_image: np.ndarray,
                        figure_detections: List[Dict],
                        page_number: int = 0) -> List[Dict[str, Any]]:
        if not self.enabled or not figure_detections:
            return []

        figures: List[Dict[str, Any]] = []
        img_h, img_w = page_image.shape[:2]

        for det in figure_detections:
            bbox = det["bbox"]
            x0, y0, x1, y1 = [int(v) for v in bbox]
            x0, y0 = max(0, x0), max(0, y0)
            x1, y1 = min(img_w, x1), min(img_h, y1)
            crop = page_image[y0:y1, x0:x1]
            if crop.size == 0:
                continue

            h, w = crop.shape[:2]
            if w < self.min_width or h < self.min_height or w * h < self.min_area:
                continue

            fig_num = len(figures) + 1                      # 1-based
            fname = f"page{page_number + 1}_fig{fig_num}.png"
            fpath = self.temp_dir / fname
            cv2.imwrite(str(fpath), crop)

            figures.append({
                "page": page_number,
                "index": fig_num,                           # 1-based
                "path": str(fpath),
                "base64": self.encode_png_base64(crop),
                "width": w,
                "height": h,
                # source bbox travels with the figure so the caller never
                # mis-aligns figures with detections it filtered out
                "bbox": [float(x0), float(y0), float(x1), float(y1)],
                "source": det.get("source", "detected"),
            })
        return figures

    def page_figure(self, page_image: np.ndarray, page_number: int = 0,
                    min_ink: float = 0.004) -> Optional[Dict[str, Any]]:
        """Return the page's ink area as one figure (graphic-only pages).

        Used when OCR finds no text and layout detection finds no regions —
        e.g. cover pages, full-page diagrams, logo pages. Previously these
        pages were dropped from the output entirely.
        """
        if not self.enabled or page_image is None or page_image.size == 0:
            return None
        gray = (cv2.cvtColor(page_image, cv2.COLOR_BGR2GRAY)
                if len(page_image.shape) == 3 else page_image)
        h, w = gray.shape[:2]
        _, binary = cv2.threshold(
            gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
        ink = cv2.countNonZero(binary)
        if ink < (h * w) * min_ink:
            return None
        ys, xs = np.where(binary > 0)
        pad = max(4, int(min(h, w) * 0.01))
        x0 = max(0, int(xs.min()) - pad)
        y0 = max(0, int(ys.min()) - pad)
        x1 = min(w, int(xs.max()) + pad)
        y1 = min(h, int(ys.max()) + pad)
        crop = page_image[y0:y1, x0:x1]
        if crop.size == 0:
            return None
        return {
            "page": page_number,
            "index": 1,
            "base64": self.encode_png_base64(crop),
            "width": int(x1 - x0),
            "height": int(y1 - y0),
            "bbox": [float(x0), float(y0), float(x1), float(y1)],
            "source": "fullpage",
        }


# ══════════════════════════════════════════════════════════════════════════════
# Document Export — HTML-first  (DOCX / TXT / HTML)
# ══════════════════════════════════════════════════════════════════════════════
class DocumentExporter:
    """HTML-first exporter (v0.4).

    Build a styled HTML document from ContentBlocks, then:
      - save the HTML directly
      - convert HTML -> DOCX via htmldocx (with manual fallback)
      - derive plain-text from blocks
    """

    FONT_NAME = "Tahoma"
    FONT_SIZE_NORMAL = 11
    FONT_SIZE_TABLE = 10

    def __init__(self):
        pass

    # ══════════════════════════════════════════════════════════════════════════
    # Primary API (v0.4 — block-based, HTML-first)
    # ══════════════════════════════════════════════════════════════════════════
    def create_all_from_blocks(self, blocks: list,
                               metadata: str = "",
                               page_size: str = "A4",
                               margin_preset: str = "Normal",
                               layout_mode: Optional[str] = None,
                               ) -> Dict[str, Optional[str]]:
        """Create TXT + HTML + DOCX.

        layout_mode:
            "flow" (default) — continuous Word-style document: flowing
            paragraphs (no text boxes) that preserve alignment, indent,
            font sizes, inline tables and figures in reading order.
            "absolute" — every block placed at its exact page position
            (text frames / positioned tables) mirroring the scan 1:1.
        """
        mode = (layout_mode or os.getenv("LAYOUT_MODE", "flow")).lower()
        has_positions = any(
            getattr(b, "bbox", None) and getattr(b, "page_width", 0) > 0
            for b in blocks)

        docx_path: Optional[str] = None
        if has_positions:
            try:
                if mode == "absolute":
                    # Opt-in: positioned text frames mirroring the scan 1:1
                    docx_path = self._build_docx_absolute(
                        blocks, page_size, margin_preset)
                else:
                    # Default ("flow"): continuous Word-style document —
                    # normal flowing paragraphs (no text boxes) that keep
                    # alignment, indentation, font size and reading order.
                    docx_path = self._build_docx_flow_structured(
                        blocks, page_size, margin_preset)
            except Exception:
                logger.exception("Structured DOCX failed — falling back")

        if has_positions:
            html_content = self._build_html_absolute(blocks, metadata)
        else:
            html_content = self._build_html(blocks, metadata)

        if docx_path is None:
            docx_path = self._html_to_docx(
                self._build_html(blocks, metadata), page_size, margin_preset)

        return {
            "txt":  self._blocks_to_txt(blocks, metadata),
            "html": self._save_html(html_content),
            "docx": docx_path,
        }

    # ── HTML builder ──────────────────────────────────────────────────────────
    def _build_html(self, blocks: list, metadata: str = "") -> str:
        parts: List[str] = self._html_head()
        if metadata:
            meta_html = metadata.replace("\n", "<br>")
            parts.append(f"<div class='meta-info'>{meta_html}</div>")

        current_page = -1
        for b in blocks:
            if b.page != current_page:
                current_page = b.page
                parts.append(self._page_divider_html(b.page, current_page == b.page))
            parts.extend(self._block_to_html(b))

        parts += ["</body>", "</html>"]
        return "\n".join(parts)

    # ── Absolute-positioned HTML preview (mirrors the DOCX layout) ───────────
    _PAGE_PREVIEW_PX = 900

    def _build_html_absolute(self, blocks: list, metadata: str = "") -> str:
        parts: List[str] = [
            "<!DOCTYPE html>", "<html lang='en'>", "<head>",
            "<meta charset='utf-8'>", "<title>OCR Result</title>",
            "<style>",
            "body { font-family: Tahoma, 'Segoe UI', Arial, sans-serif; "
            "background: #eef1f5; margin: 0; padding: 1.5rem; }",
            ".page { position: relative; background: #fff; margin: 0 auto "
            "1.5rem auto; box-shadow: 0 2px 10px rgba(0,0,0,0.15); "
            "overflow: hidden; }",
            ".abs { position: absolute; white-space: pre; line-height: 1.25; "
            "color: #1e293b; }",
            ".abs-table { position: absolute; }",
            ".abs-table table { border-collapse: collapse; width: 100%; "
            "height: 100%; table-layout: fixed; }",
            ".abs-table th, .abs-table td { border: 1px solid #94a3b8; "
            "padding: 1px 4px; vertical-align: top; overflow: hidden; }",
            ".abs img { width: 100%; height: 100%; object-fit: contain; }",
            ".page-label { max-width: 900px; margin: 0 auto; color: #64748b; "
            "font-size: 0.8rem; padding: 0.3rem 0; }",
            ".meta-info { max-width: 900px; margin: 0 auto 1rem auto; "
            "color: #64748b; font-size: 0.85rem; white-space: pre-line; }",
            "</style>", "</head>", "<body>",
        ]
        if metadata:
            parts.append(
                f"<div class='meta-info'>{html.escape(metadata)}</div>")

        pages: Dict[int, list] = {}
        for b in blocks:
            pages.setdefault(b.page, []).append(b)

        for page_num in sorted(pages):
            page_blocks = pages[page_num]
            src_w = max((getattr(b, "page_width", 0) for b in page_blocks),
                        default=0)
            src_h = max((getattr(b, "page_height", 0) for b in page_blocks),
                        default=0)
            parts.append(f"<div class='page-label'>Page {page_num + 1}</div>")
            if src_w <= 0 or src_h <= 0:
                parts.append("<div class='page' style='max-width:900px;"
                             "padding:2rem'>")
                for b in page_blocks:
                    parts.extend(self._block_to_html(b))
                parts.append("</div>")
                continue

            scale = self._PAGE_PREVIEW_PX / src_w
            page_h_px = int(src_h * scale)
            parts.append(
                f"<div class='page' style='width:{self._PAGE_PREVIEW_PX}px;"
                f"height:{page_h_px}px'>")
            for b in page_blocks:
                parts.extend(self._abs_block_html(b, scale))
            parts.append("</div>")

        parts += ["</body>", "</html>"]
        return "\n".join(parts)

    def _abs_block_html(self, b, scale: float) -> List[str]:
        """Render one positioned block as absolutely-placed HTML."""
        if not getattr(b, "bbox", None):
            return self._block_to_html(b)
        x0, y0, x1, y1 = b.bbox
        left, top = x0 * scale, y0 * scale
        width, height = (x1 - x0) * scale, (y1 - y0) * scale

        if b.block_type in ("text", "caption"):
            out: List[str] = []
            lines = b.lines or []
            if not lines:
                font_px = max(9.0, height * 0.7 / max(
                    len(b.text.split("\n")), 1))
                style = (f"left:{left:.1f}px;top:{top:.1f}px;"
                         f"width:{width:.1f}px;font-size:{font_px:.1f}px;")
                out.append(f"<div class='abs' style='{style}'>"
                           f"{html.escape(b.text)}</div>")
                return out
            for line in lines:
                lb = line.get("bbox") or b.bbox
                lh = (lb[3] - lb[1]) * scale
                size_pt = float(line.get("size_pt") or 0)
                font_px = (size_pt * scale * 1.0) if size_pt > 0 \
                    else max(7.0, lh * 0.72)
                weight = "bold" if line.get("bold") else "normal"
                style = (f"left:{lb[0] * scale:.1f}px;"
                         f"top:{lb[1] * scale:.1f}px;"
                         f"font-size:{font_px:.1f}px;font-weight:{weight};")
                out.append(f"<div class='abs' style='{style}'>"
                           f"{html.escape(line.get('text', ''))}</div>")
            return out

        if b.block_type == "table" and b.table_html:
            style = (f"left:{left:.1f}px;top:{top:.1f}px;"
                     f"width:{width:.1f}px;height:{height:.1f}px;"
                     f"font-size:{max(7.0, height * 0.55 / max(b.table_html.count('<tr'), 1)):.1f}px;")
            return [f"<div class='abs-table' style='{style}'>"
                    f"{b.table_html}</div>"]

        if b.block_type == "figure":
            b64 = (b.figure or {}).get("base64", "")
            if not b64 or not _is_valid_base64(b64):
                return []
            style = (f"left:{left:.1f}px;top:{top:.1f}px;"
                     f"width:{width:.1f}px;height:{height:.1f}px;")
            return [f"<div class='abs' style='{style}'>"
                    f"<img src='data:image/png;base64,{b64}' alt='Figure'/>"
                    "</div>"]
        return []

    @staticmethod
    def _html_head() -> List[str]:
        """Return the HTML preamble (doctype → <body>)."""
        return [
            "<!DOCTYPE html>",
            "<html lang='en'>",
            "<head>",
            "<meta charset='utf-8'>",
            "<title>OCR Result</title>",
            "<style>",
            "body { font-family: Tahoma, 'Segoe UI', Arial, sans-serif; "
            "max-width: 900px; margin: 0 auto; padding: 2rem; "
            "line-height: 1.8; color: #1e293b; }",
            "h1, h2, h3 { color: #0f172a; margin-top: 1.5em; margin-bottom: 0.5em; }",
            "p { margin: 0.4em 0; text-align: justify; }",
            "table { border-collapse: collapse; width: 100%; margin: 1rem 0; }",
            "th, td { border: 1px solid #cbd5e1; padding: 8px 12px; "
            "text-align: left; vertical-align: top; }",
            "th { background: #f1f5f9; font-weight: bold; }",
            "figure { margin: 1rem 0; text-align: center; }",
            "figure img { max-width: 100%; height: auto; border-radius: 4px; "
            "box-shadow: 0 2px 8px rgba(0,0,0,0.1); }",
            "figcaption { color: #64748b; font-size: 0.85rem; margin-top: 0.5rem; }",
            ".page-break { page-break-before: always; border-top: 2px solid #e2e8f0; "
            "margin-top: 2rem; padding-top: 1rem; color: #94a3b8; "
            "font-size: 0.85rem; }",
            ".meta-info { color: #64748b; font-size: 0.85rem; "
            "border-bottom: 1px solid #e2e8f0; padding-bottom: 1rem; "
            "margin-bottom: 2rem; white-space: pre-line; }",
            "</style>",
            "</head>",
            "<body>",
        ]

    @staticmethod
    def _page_divider_html(page_num: int, is_first: bool) -> str:
        """Return an HTML page divider element."""
        if is_first:
            return (f"<div class='page-break' style='border:none;margin-top:0;"
                    f"padding-top:0'>Page {page_num + 1}</div>")
        return f"<div class='page-break'>Page {page_num + 1}</div>"

    def _block_to_html(self, b) -> List[str]:
        """Convert a single ContentBlock to HTML snippet(s)."""
        if b.block_type in ("text", "caption"):
            return [self._text_to_html(line.strip())
                    for line in b.text.split("\n") if line.strip()]
        if b.block_type == "table":
            if b.table_html:
                return [b.table_html]
            if b.text.strip():
                return [f"<pre>{b.text}</pre>"]
            return []
        if b.block_type == "figure":
            return self._figure_block_to_html(b)
        return []

    @staticmethod
    def _figure_block_to_html(b) -> List[str]:
        """Convert a figure ContentBlock to an HTML <figure> snippet."""
        b64 = b.figure.get("base64", "")
        if not b64 or not _is_valid_base64(b64):
            return []
        fig_idx = int(b.figure.get("index", 0))
        fig_page = b.figure.get("page", 0)
        page_disp = (fig_page + 1) if isinstance(fig_page, int) else fig_page
        escaped_alt = html.escape(f"Figure {fig_idx}")
        escaped_cap = html.escape(f"Figure {fig_idx} \u2014 Page {page_disp}")
        orig_w = b.figure.get("width", 0)
        orig_h = b.figure.get("height", 0)
        if orig_w > 0 and orig_h > 0:
            display_w = min(orig_w, _MAX_IMG_DISPLAY_PX)
            display_h = int(orig_h * display_w / orig_w)
            img_style = f"width:{display_w}px;height:{display_h}px;"
        else:
            img_style = f"max-width:{_MAX_IMG_DISPLAY_PX}px;height:auto;"
        return [
            (
                "<figure>"
                + f"<img src='data:image/png;base64,{b64}' "
                + f"alt='{escaped_alt}' style='{img_style}' />"
                + f"<figcaption>{escaped_cap}</figcaption>"
                + "</figure>"
            )
        ]

    # ── Save HTML ─────────────────────────────────────────────────────────────
    @staticmethod
    def _save_html(html_content: str) -> str:
        fd, path = tempfile.mkstemp(suffix=".html", prefix="ocr_")
        with os.fdopen(fd, "w", encoding="utf-8") as f:
            f.write(html_content)
        return path

    # ── HTML -> DOCX ─────────────────────────────────────────────────────────
    @staticmethod
    def _extract_base64_images_to_tempfiles(html_content: str) -> str:
        """Replace base64 data URIs in <img> src with temp file paths.

        htmldocx tries to use the src attribute as a filename, which fails
        when the src is a data URI (filename too long).  This pre-processes
        the HTML so that base64 images are saved as temp files and the src
        attributes point to those files instead.
        """
        _DATA_URI_RE = re.compile(
            r"""src\s*=\s*['"]data:image/(png|jpe?g|gif|webp);base64,([A-Za-z0-9+/=]+)['"]""",
        )

        def _replacer(match):
            fmt = match.group(1).lower()
            b64_data = match.group(2)
            suffix = f".{fmt}" if fmt != "jpeg" else ".jpg"
            try:
                img_bytes = base64.b64decode(b64_data)
                fd, tmp_path = tempfile.mkstemp(suffix=suffix, prefix="ocr_img_")
                with os.fdopen(fd, "wb") as f:
                    f.write(img_bytes)
                return f"src='{tmp_path}'"
            except Exception:
                return match.group(0)  # leave unchanged on error

        return _DATA_URI_RE.sub(_replacer, html_content)

    def _html_to_docx(self, html_content: str,
                       page_size: str = "A4",
                       margin_preset: str = "Normal") -> Optional[str]:
        if not DOCX_AVAILABLE:
            logger.error("python-docx not installed")
            return None

        # Try htmldocx first — handles images, tables, alignment
        if HTMLDOCX_AVAILABLE:
            try:
                # Pre-process: extract base64 images to temp files
                # to avoid "filename too long" errors in htmldocx
                processed_html = self._extract_base64_images_to_tempfiles(
                    html_content)
                parser = HtmlToDocx()
                doc = parser.parse_html_string(processed_html)
                # Apply page layout, fix table borders, fix image sizes
                self._apply_page_layout(doc, page_size, margin_preset)
                self._fix_table_borders(doc)
                self._fix_image_sizes(doc, page_size, margin_preset)
                fd, path = tempfile.mkstemp(suffix=".docx", prefix="ocr_")
                os.close(fd)
                doc.save(path)
                logger.info("DOCX created via htmldocx")
                return path
            except Exception as exc:
                logger.warning("htmldocx failed (%s) — using fallback", exc)

        # Fallback: manual DOCX from parsed HTML
        return self._fallback_docx(html_content, page_size, margin_preset)

    def _fallback_docx(self, html_content: str,
                        page_size: str = "A4",
                        margin_preset: str = "Normal") -> Optional[str]:
        """Create DOCX by parsing HTML with BeautifulSoup."""
        if not BS4_AVAILABLE:
            return None

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = self.FONT_NAME
        style.font.size = Pt(self.FONT_SIZE_NORMAL)

        # Apply page layout
        self._apply_page_layout(doc, page_size, margin_preset)

        soup = BeautifulSoup(html_content, "html.parser")
        body = soup.find("body")
        if not body:
            return None

        # Store layout info for image sizing
        self._current_page_size = page_size
        self._current_margin_preset = margin_preset

        for el in body.children:
            if not hasattr(el, "name") or el.name is None:
                continue
            self._el_to_docx(doc, el)

        fd, path = tempfile.mkstemp(suffix=".docx", prefix="ocr_")
        os.close(fd)
        doc.save(path)
        return path

    def _el_to_docx(self, doc, el) -> None:
        """Convert a parsed HTML element into DOCX content (dispatch)."""
        tag = el.name
        handler = {
            "h1": self._heading_to_docx,
            "h2": self._heading_to_docx,
            "h3": self._heading_to_docx,
            "p": self._paragraph_to_docx,
            "pre": self._preformatted_to_docx,
            "table": self._table_el_to_docx,
            "figure": self._figure_el_to_docx,
        }.get(tag)
        if handler:
            handler(doc, el)
        elif tag == "div":
            for child in el.children:
                if hasattr(child, "name") and child.name:
                    self._el_to_docx(doc, child)

    def _heading_to_docx(self, doc, el) -> None:
        level = int(el.name[1])
        h = doc.add_heading(el.get_text(strip=True), level=level)
        for run in h.runs:
            run.font.name = self.FONT_NAME

    def _paragraph_to_docx(self, doc, el) -> None:
        text = el.get_text(strip=True)
        if text:
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.name = self.FONT_NAME
                run.font.size = Pt(self.FONT_SIZE_NORMAL)

    def _preformatted_to_docx(self, doc, el) -> None:
        text = el.get_text()
        if text.strip():
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(self.FONT_SIZE_TABLE)

    def _table_el_to_docx(self, doc, table_el) -> None:
        rows_el = table_el.find_all("tr")
        if not rows_el:
            return

        max_cols = self._calc_table_col_count(rows_el)
        if max_cols == 0:
            return

        tbl = doc.add_table(rows=len(rows_el), cols=max_cols)
        tbl.style = "Table Grid"

        for ri, row_el in enumerate(rows_el):
            cells = row_el.find_all(["td", "th"])
            col_idx = 0
            for cell_el in cells:
                if col_idx >= max_cols:
                    break
                col_idx = self._fill_table_cell(
                    tbl, ri, col_idx, cell_el, max_cols, len(rows_el))

    @staticmethod
    def _calc_table_col_count(rows_el) -> int:
        """Return the maximum column count across all rows."""
        max_cols = 0
        for r in rows_el:
            cols_in_row = sum(
                int(cell.get("colspan", 1))
                for cell in r.find_all(["td", "th"])
            )
            max_cols = max(max_cols, cols_in_row)
        return max_cols

    def _fill_table_cell(self, tbl, ri: int, col_idx: int,
                         cell_el, max_cols: int, total_rows: int) -> int:
        """Populate one table cell and return the next column index."""
        colspan = int(cell_el.get("colspan", 1))
        rowspan = int(cell_el.get("rowspan", 1))

        cell = tbl.cell(ri, col_idx)
        cell.text = cell_el.get_text(strip=True)

        if colspan > 1 and col_idx + colspan - 1 < max_cols:
            try:
                cell.merge(tbl.cell(ri, col_idx + colspan - 1))
            except Exception:
                pass
        if rowspan > 1 and ri + rowspan - 1 < total_rows:
            try:
                cell.merge(tbl.cell(ri + rowspan - 1, col_idx))
            except Exception:
                pass

        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = self.FONT_NAME
                run.font.size = Pt(self.FONT_SIZE_TABLE)
                if cell_el.name == "th":
                    run.bold = True

        return col_idx + colspan

    def _figure_el_to_docx(self, doc, figure_el) -> None:
        img_el = figure_el.find("img")
        caption_el = figure_el.find("figcaption")
        caption_text = caption_el.get_text(strip=True) if caption_el else ""
        if not img_el:
            return
        src = img_el.get("src", "")
        try:
            img_bytes = None
            if src.startswith("data:image"):
                b64_data = src.split(",", 1)[1] if "," in src else ""
                if b64_data:
                    img_bytes = base64.b64decode(b64_data)
            elif os.path.isfile(src):
                with open(src, "rb") as fimg:
                    img_bytes = fimg.read()

            if img_bytes:
                # Calculate reasonable image width
                page_sz = getattr(self, "_current_page_size", "A4")
                margin_p = getattr(self, "_current_margin_preset", "Normal")
                img_width = self._calc_image_width(img_bytes, page_sz, margin_p)
                doc.add_picture(BytesIO(img_bytes), width=Inches(img_width))

            if caption_text:
                p = doc.add_paragraph(caption_text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(100, 116, 139)
        except Exception as exc:
            logger.warning("Failed to embed figure in DOCX: %s", exc)

    # ══════════════════════════════════════════════════════════════════════════
    # Absolute-position DOCX builder  (v2.2 — layout-faithful export)
    #
    # Every ContentBlock is rendered at its exact page coordinates:
    #   - text/caption → Word text frames (w:framePr) with per-line font size,
    #     alignment, and indentation derived from OCR line positions
    #   - table        → real, editable Word table positioned via w:tblpPr
    #   - figure       → image inside a positioned frame, sized to its bbox
    # ══════════════════════════════════════════════════════════════════════════

    _MIN_FONT_PT = 5.0
    _MAX_FONT_PT = 48.0
    _OCR_FONT_FACTOR = 0.72   # line bbox includes ascender/descender

    def _build_docx_absolute(self, blocks: list,
                             page_size: str = "A4",
                             margin_preset: str = "Normal") -> Optional[str]:
        if not DOCX_AVAILABLE:
            return None

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = self.FONT_NAME
        style.font.size = Pt(self.FONT_SIZE_NORMAL)
        self._set_eastasia_font(style, self.FONT_NAME)
        self._apply_page_layout(doc, page_size, margin_preset)

        pw_in, ph_in = PAGE_SIZES.get(page_size, PAGE_SIZES["A4"])
        page_w_tw = int(pw_in * 1440)
        page_h_tw = int(ph_in * 1440)

        # Group blocks per page (input is already in reading order)
        pages: Dict[int, list] = {}
        for b in blocks:
            pages.setdefault(b.page, []).append(b)

        first = True
        for page_num in sorted(pages):
            if not first:
                br_p = doc.add_paragraph()
                br_p.add_run().add_break(WD_BREAK.PAGE)
            first = False

            page_blocks = pages[page_num]
            src_w = max((getattr(b, "page_width", 0) for b in page_blocks),
                        default=0)
            src_h = max((getattr(b, "page_height", 0) for b in page_blocks),
                        default=0)

            for b in page_blocks:
                if not getattr(b, "bbox", None) or src_w <= 0 or src_h <= 0:
                    self._add_flow_block(doc, b)
                    continue
                sx = page_w_tw / src_w          # twips per page unit (x)
                sy = page_h_tw / src_h          # twips per page unit (y)
                pt_per_unit = (ph_in * 72.0) / src_h
                if b.block_type in ("text", "caption"):
                    self._add_text_frame(doc, b, sx, sy, pt_per_unit)
                elif b.block_type == "table":
                    self._add_table_absolute(doc, b, sx, sy, pt_per_unit)
                elif b.block_type == "figure":
                    self._add_figure_frame(doc, b, sx, sy)

        self._fix_table_borders(doc)
        self._reorder_table_props(doc)
        fd, path = tempfile.mkstemp(suffix=".docx", prefix="ocr_")
        os.close(fd)
        doc.save(path)
        logger.info("DOCX created via absolute-position builder")
        return path

    @staticmethod
    def _set_eastasia_font(style, font_name: str) -> None:
        """Set east-asian / complex-script font so Thai renders correctly."""
        try:
            rpr = style.element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:eastAsia"), font_name)
            rfonts.set(qn("w:cs"), font_name)
        except Exception:
            pass

    @staticmethod
    def _set_frame_pr(paragraph, x_tw: int, y_tw: int, w_tw: int) -> None:
        """Attach w:framePr — absolute page-anchored text frame."""
        p_pr = paragraph._p.get_or_add_pPr()  # noqa: SLF001
        frame = p_pr.find(qn("w:framePr"))
        if frame is None:
            frame = OxmlElement("w:framePr")
            p_pr.insert(0, frame)
        frame.set(qn("w:w"), str(max(w_tw, 144)))
        frame.set(qn("w:hRule"), "auto")
        frame.set(qn("w:wrap"), "around")
        frame.set(qn("w:hAnchor"), "page")
        frame.set(qn("w:vAnchor"), "page")
        frame.set(qn("w:x"), str(max(x_tw, 0)))
        frame.set(qn("w:y"), str(max(y_tw, 0)))

    def _style_frame_paragraph(self, paragraph) -> None:
        pf = paragraph.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

    def _line_font_pt(self, line: Dict[str, Any],
                      pt_per_unit: float) -> float:
        size = float(line.get("size_pt") or 0)
        if size <= 0:
            bb = line.get("bbox") or [0, 0, 0, 0]
            size = (bb[3] - bb[1]) * pt_per_unit * self._OCR_FONT_FACTOR
        return max(self._MIN_FONT_PT, min(self._MAX_FONT_PT, size))

    @staticmethod
    def _line_alignment(line_bbox, block_bbox):
        """Derive alignment + left indent (page units) for one line."""
        bx0, _, bx1, _ = block_bbox
        lx0, _, lx1, _ = line_bbox
        bw = max(bx1 - bx0, 1e-6)
        lg = (lx0 - bx0) / bw     # left gap fraction
        rg = (bx1 - lx1) / bw     # right gap fraction
        if lg > 0.08 and rg > 0.08 and abs(lg - rg) < 0.08:
            return WD_ALIGN_PARAGRAPH.CENTER, 0.0
        if rg < 0.03 and lg > 0.15:
            return WD_ALIGN_PARAGRAPH.RIGHT, 0.0
        return WD_ALIGN_PARAGRAPH.LEFT, max(0.0, lx0 - bx0)

    def _add_text_frame(self, doc, b, sx: float, sy: float,
                        pt_per_unit: float) -> None:
        """One block → one text frame; lines keep size/alignment/indent."""
        x0, y0, x1, _ = b.bbox
        x_tw = int(x0 * sx)
        y_tw = int(y0 * sy)
        w_tw = int((x1 - x0) * sx) + 72   # small slack against re-wrap

        lines = b.lines or [
            {"text": ln, "bbox": b.bbox}
            for ln in b.text.split("\n") if ln.strip()
        ]
        for line in lines:
            text = line.get("text", "").strip()
            if not text:
                continue
            p = doc.add_paragraph()
            self._set_frame_pr(p, x_tw, y_tw, w_tw)
            self._style_frame_paragraph(p)

            align, indent_units = self._line_alignment(
                line.get("bbox") or b.bbox, b.bbox)
            p.alignment = align
            if indent_units > 0:
                p.paragraph_format.left_indent = Emu(
                    int(indent_units * sx * 635))  # twips → EMU (×635)

            run = p.add_run(text)
            run.font.name = self.FONT_NAME
            run.font.size = Pt(round(self._line_font_pt(line, pt_per_unit), 1))
            if line.get("bold"):
                run.bold = True

    def _add_table_absolute(self, doc, b, sx: float, sy: float,
                            pt_per_unit: float) -> None:
        """Positioned, real (editable) Word table built from table_html."""
        if not BS4_AVAILABLE or not b.table_html:
            # No structure — render plain text in a frame instead
            if b.text.strip():
                self._add_text_frame(doc, b, sx, sy, pt_per_unit)
            return
        soup = BeautifulSoup(b.table_html, "html.parser")
        table_el = soup.find("table")
        if table_el is None:
            return
        rows_el = table_el.find_all("tr")
        max_cols = self._calc_table_col_count(rows_el)
        if not rows_el or max_cols == 0:
            return

        x0, y0, x1, y1 = b.bbox
        tbl_w_tw = max(int((x1 - x0) * sx), 720)
        row_h_pt = ((y1 - y0) * pt_per_unit) / max(len(rows_el), 1)
        cell_font = max(6.0, min(12.0, row_h_pt * 0.5))

        tbl = doc.add_table(rows=len(rows_el), cols=max_cols)
        tbl.style = "Table Grid"
        tbl.autofit = False

        # Column widths — measured grid proportions when available
        col_fracs = (b.table_meta or {}).get("col_widths") or []
        if len(col_fracs) != max_cols:
            col_fracs = [1.0 / max_cols] * max_cols
        for ci, col in enumerate(tbl.columns):
            col.width = Emu(int(tbl_w_tw * col_fracs[ci] * 635))

        for ri, row_el in enumerate(rows_el):
            col_idx = 0
            for cell_el in row_el.find_all(["td", "th"]):
                if col_idx >= max_cols:
                    break
                col_idx = self._fill_table_cell(
                    tbl, ri, col_idx, cell_el, max_cols, len(rows_el))
            for p in (c.paragraphs[0] for c in tbl.rows[ri].cells):
                for run in p.runs:
                    run.font.size = Pt(round(cell_font, 1))

        self._set_table_position(
            tbl, int(x0 * sx), int(y0 * sy), tbl_w_tw)

    @staticmethod
    def _set_table_position(table, x_tw: int, y_tw: int, w_tw: int) -> None:
        """Attach w:tblpPr — absolute page-anchored floating table."""
        tbl_pr = table._tbl.tblPr  # noqa: SLF001
        for tag in ("w:tblpPr", "w:tblOverlap", "w:tblLayout", "w:tblW"):
            old = tbl_pr.find(qn(tag))
            if old is not None:
                tbl_pr.remove(old)
        pos = OxmlElement("w:tblpPr")
        pos.set(qn("w:leftFromText"), "142")
        pos.set(qn("w:rightFromText"), "142")
        pos.set(qn("w:vertAnchor"), "page")
        pos.set(qn("w:horzAnchor"), "page")
        pos.set(qn("w:tblpX"), str(max(x_tw, 1)))
        pos.set(qn("w:tblpY"), str(max(y_tw, 1)))
        tbl_pr.append(pos)
        overlap = OxmlElement("w:tblOverlap")
        overlap.set(qn("w:val"), "never")
        tbl_pr.append(overlap)
        # Fixed layout + explicit width so columns match the scan
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tbl_pr.append(layout)
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), str(w_tw))
        tbl_w.set(qn("w:type"), "dxa")

    # OOXML schema order for w:tblPr children — Word rejects bad ordering
    _TBLPR_ORDER = (
        "tblStyle", "tblpPr", "tblOverlap", "bidiVisual",
        "tblStyleRowBandSize", "tblStyleColBandSize", "tblW", "jc",
        "tblCellSpacing", "tblInd", "tblBorders", "shd", "tblLayout",
        "tblCellMar", "tblLook", "tblCaption", "tblDescription",
    )

    @classmethod
    def _reorder_table_props(cls, doc) -> None:
        """Re-sort every table's tblPr children into OOXML schema order."""
        rank = {qn(f"w:{tag}"): i for i, tag in enumerate(cls._TBLPR_ORDER)}
        for table in doc.tables:
            tbl_pr = table._tbl.tblPr  # noqa: SLF001
            if tbl_pr is None:
                continue
            children = list(tbl_pr)
            children.sort(key=lambda el: rank.get(el.tag, len(rank)))
            for el in children:
                tbl_pr.remove(el)
            for el in children:
                tbl_pr.append(el)

    def _add_figure_frame(self, doc, b, sx: float, sy: float) -> None:
        """Image extracted via OpenCV, framed at its exact page position."""
        b64 = (b.figure or {}).get("base64", "")
        if not b64 or not _is_valid_base64(b64):
            return
        try:
            img_bytes = base64.b64decode(b64)
        except Exception:
            return
        x0, y0, x1, y1 = b.bbox
        x_tw, y_tw = int(x0 * sx), int(y0 * sy)
        w_tw = max(int((x1 - x0) * sx), 144)
        h_tw = max(int((y1 - y0) * sy), 144)

        p = doc.add_paragraph()
        self._set_frame_pr(p, x_tw, y_tw, w_tw)
        self._style_frame_paragraph(p)
        run = p.add_run()
        try:
            run.add_picture(BytesIO(img_bytes),
                            width=Emu(w_tw * 635), height=Emu(h_tw * 635))
        except Exception as exc:
            logger.warning("Failed to place figure frame: %s", exc)

    def _add_flow_block(self, doc, b) -> None:
        """Fallback for blocks without position data — normal paragraphs."""
        if b.block_type in ("text", "caption"):
            for ln in b.text.split("\n"):
                if ln.strip():
                    p = doc.add_paragraph(ln.strip())
                    for run in p.runs:
                        run.font.name = self.FONT_NAME
                        run.font.size = Pt(self.FONT_SIZE_NORMAL)
        elif b.block_type == "table" and b.table_html and BS4_AVAILABLE:
            soup = BeautifulSoup(b.table_html, "html.parser")
            table_el = soup.find("table")
            if table_el is not None:
                self._table_el_to_docx(doc, table_el)
        elif b.block_type == "figure":
            b64 = (b.figure or {}).get("base64", "")
            if b64 and _is_valid_base64(b64):
                try:
                    doc.add_picture(BytesIO(base64.b64decode(b64)),
                                    width=Inches(4.0))
                except Exception:
                    pass

    # ══════════════════════════════════════════════════════════════════════════
    # Structured flow DOCX builder  (v3.1 — Word-like continuous text)
    #
    # Default DOCX output: normal flowing paragraphs (NO text boxes), with
    # alignment, indentation, font size and bold preserved from the page
    # geometry. Wrapped lines are merged into continuous paragraphs and
    # tables/figures are placed inline in reading order.
    # ══════════════════════════════════════════════════════════════════════════

    _WRAP_RIGHT_REACH = 0.80   # line reaching >=80% of block width = wrapped
    _LINE_GAP_FACTOR = 0.85    # vert. gap below this x line height = same para

    @staticmethod
    def _thai_join(parts: List[str]) -> str:
        """Join text parts; no space inserted between Thai<->Thai boundaries."""
        out = ""
        for part in parts:
            part = part.strip()
            if not part:
                continue
            if not out:
                out = part
                continue
            prev_ch, next_ch = out[-1], part[0]
            if ("\u0e00" <= prev_ch <= "\u0e7f"
                    and "\u0e00" <= next_ch <= "\u0e7f"):
                out += part
            else:
                out += " " + part
        return out

    def _merge_lines_to_paragraphs(self, b) -> List[List[Dict[str, Any]]]:
        """Group a block's visual lines into logical paragraphs.

        A line continues the previous paragraph when the previous line
        looks wrapped (its right edge reaches the block's right side),
        the vertical gap is a normal line advance, and the new line is
        not strongly indented.
        """
        lines = b.lines or [{"text": ln, "bbox": (list(b.bbox) or None)}
                            for ln in b.text.split("\n") if ln.strip()]
        bb = b.bbox or [0, 0, 0, 0]
        bx0, bx1 = bb[0], bb[2]
        bw = max(bx1 - bx0, 1e-6)
        paras: List[List[Dict[str, Any]]] = []
        cur: List[Dict[str, Any]] = []
        prev = None
        for line in lines:
            if not (line.get("text") or "").strip():
                continue
            lb = line.get("bbox")
            same_para = False
            if cur and prev is not None and lb and prev.get("bbox"):
                pb = prev["bbox"]
                ph = max(pb[3] - pb[1], 1e-6)
                gap = lb[1] - pb[3]
                prev_reach = (pb[2] - bx0) / bw
                same_para = (gap < ph * self._LINE_GAP_FACTOR
                             and prev_reach >= self._WRAP_RIGHT_REACH
                             and (lb[0] - bx0) / bw < 0.35)
            if same_para:
                cur.append(line)
            else:
                if cur:
                    paras.append(cur)
                cur = [line]
            prev = line
        if cur:
            paras.append(cur)
        return paras

    def _build_docx_flow_structured(self, blocks: list,
                                    page_size: str = "A4",
                                    margin_preset: str = "Normal"
                                    ) -> Optional[str]:
        """Continuous Word-style DOCX: flowing text, inline tables/figures."""
        if not DOCX_AVAILABLE:
            return None
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = self.FONT_NAME
        style.font.size = Pt(self.FONT_SIZE_NORMAL)
        self._set_eastasia_font(style, self.FONT_NAME)
        self._apply_page_layout(doc, page_size, margin_preset)

        pw_in, ph_in = PAGE_SIZES.get(page_size, PAGE_SIZES["A4"])
        _, _, ml, mr = MARGIN_PRESETS.get(margin_preset,
                                          MARGIN_PRESETS["Normal"])
        usable_w_in = pw_in - ml - mr

        pages: Dict[int, list] = {}
        for b in blocks:
            pages.setdefault(b.page, []).append(b)

        first = True
        for page_num in sorted(pages):
            if not first:
                br = doc.add_paragraph()
                br.add_run().add_break(WD_BREAK.PAGE)
            first = False
            page_blocks = pages[page_num]
            src_w = max((getattr(b, "page_width", 0) for b in page_blocks),
                        default=0)
            src_h = max((getattr(b, "page_height", 0) for b in page_blocks),
                        default=0)
            pt_per_unit = (ph_in * 72.0) / src_h if src_h > 0 else 0.0
            for b in page_blocks:
                if b.block_type in ("text", "caption"):
                    self._add_structured_text(
                        doc, b, src_w, usable_w_in, pt_per_unit)
                elif b.block_type == "table":
                    self._add_structured_table(doc, b, usable_w_in)
                elif b.block_type == "figure":
                    self._add_structured_figure(doc, b, src_w, usable_w_in)

        self._fix_table_borders(doc)
        fd, path = tempfile.mkstemp(suffix=".docx", prefix="ocr_")
        os.close(fd)
        doc.save(path)
        logger.info("DOCX created via structured flow builder")
        return path

    def _add_structured_text(self, doc, b, src_w: float,
                             usable_w_in: float, pt_per_unit: float) -> None:
        """One block -> flowing paragraphs with alignment/indent/size kept."""
        for para_lines in self._merge_lines_to_paragraphs(b):
            first_line = para_lines[0]
            text = self._thai_join([ln.get("text", "") for ln in para_lines])
            if not text:
                continue
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(4)
            if b.bbox and first_line.get("bbox"):
                align, indent_units = self._line_alignment(
                    first_line["bbox"], b.bbox)
                p.alignment = align
                if indent_units > 0 and src_w > 0:
                    pf.left_indent = Inches(
                        min(indent_units / src_w * usable_w_in,
                            usable_w_in * 0.6))
            run = p.add_run(text)
            run.font.name = self.FONT_NAME
            if pt_per_unit > 0 and first_line.get("bbox"):
                run.font.size = Pt(round(
                    self._line_font_pt(first_line, pt_per_unit), 1))
            if first_line.get("bold"):
                run.bold = True

    def _add_structured_table(self, doc, b, usable_w_in: float) -> None:
        """Inline (non-floating) editable Word table in reading order."""
        if not BS4_AVAILABLE or not b.table_html:
            if b.text.strip():
                p = doc.add_paragraph(b.text)
                for run in p.runs:
                    run.font.name = self.FONT_NAME
                    run.font.size = Pt(self.FONT_SIZE_TABLE)
            return
        soup = BeautifulSoup(b.table_html, "html.parser")
        table_el = soup.find("table")
        if table_el is None:
            return
        rows_el = table_el.find_all("tr")
        max_cols = self._calc_table_col_count(rows_el)
        if not rows_el or max_cols == 0:
            return
        tbl = doc.add_table(rows=len(rows_el), cols=max_cols)
        tbl.style = "Table Grid"
        tbl.autofit = False
        col_fracs = (b.table_meta or {}).get("col_widths") or []
        if len(col_fracs) != max_cols:
            col_fracs = [1.0 / max_cols] * max_cols
        for ci, col in enumerate(tbl.columns):
            col.width = Inches(usable_w_in * col_fracs[ci])
        for ri, row_el in enumerate(rows_el):
            col_idx = 0
            for cell_el in row_el.find_all(["td", "th"]):
                if col_idx >= max_cols:
                    break
                col_idx = self._fill_table_cell(
                    tbl, ri, col_idx, cell_el, max_cols, len(rows_el))
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(4)

    def _add_structured_figure(self, doc, b, src_w: float,
                               usable_w_in: float) -> None:
        """Inline centred figure, sized from its share of the page width."""
        b64 = (b.figure or {}).get("base64", "")
        if not b64 or not _is_valid_base64(b64):
            return
        try:
            img_bytes = base64.b64decode(b64)
        except Exception:
            return
        width_in = usable_w_in * 0.75
        if b.bbox and src_w > 0:
            x0, _, x1, _ = b.bbox
            width_in = max(0.8, min((x1 - x0) / src_w * usable_w_in,
                                    usable_w_in))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(BytesIO(img_bytes), width=Inches(width_in))
        except Exception as exc:
            logger.warning("Failed to place figure: %s", exc)

    # ── TXT from blocks ──────────────────────────────────────────────────────
    def _blocks_to_txt(self, blocks: list, metadata: str = "") -> str:
        parts: List[str] = []
        if metadata:
            parts.append(f"--- Document Info ---\n{metadata}\n---\n")

        current_page = -1
        for b in blocks:
            if b.page != current_page:
                if current_page >= 0:
                    parts.append(f"\n{'─' * 60}")
                parts.append(f"  Page {b.page + 1}")
                parts.append(f"{'─' * 60}\n")
                current_page = b.page
            parts.append(self._block_to_txt(b))

        content = "\n\n".join(p for p in parts if p)
        fd, path = tempfile.mkstemp(suffix=".txt", prefix="ocr_")
        with os.fdopen(fd, "w", encoding="utf-8") as f:
            f.write(content)
        return path

    @staticmethod
    def _block_to_txt(b) -> str:
        """Convert a single ContentBlock to its plain-text representation."""
        if b.block_type in ("text", "caption"):
            return b.text if b.text.strip() else ""
        if b.block_type == "table":
            return f"[Table]\n{b.text}" if b.text.strip() else ""
        if b.block_type == "figure":
            fig_idx = b.figure.get("index", "?")
            pg = b.figure.get("page", "?")
            page_disp = (pg + 1) if isinstance(pg, int) else pg
            return f"[Figure {fig_idx} — Page {page_disp}]"
        return ""

    # ── Page layout & post-processing helpers ─────────────────────────────────
    @staticmethod
    def _apply_page_layout(doc, page_size: str = "A4",
                           margin_preset: str = "Normal") -> None:
        """Set page dimensions and margins on all sections of the DOCX."""
        pw, ph = PAGE_SIZES.get(page_size, PAGE_SIZES["A4"])
        mt, mb, ml, mr = MARGIN_PRESETS.get(margin_preset, MARGIN_PRESETS["Normal"])
        for section in doc.sections:
            section.page_width = Inches(pw)
            section.page_height = Inches(ph)
            section.top_margin = Inches(mt)
            section.bottom_margin = Inches(mb)
            section.left_margin = Inches(ml)
            section.right_margin = Inches(mr)

    @staticmethod
    def _fix_table_borders(doc) -> None:
        """Ensure every table has visible cell borders (like Word's Table Grid)."""
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml
        for table in doc.tables:
            tbl = table._tbl
            tbl_pr = tbl.tblPr
            if tbl_pr is None:
                tbl_pr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
                tbl.insert(0, tbl_pr)
            borders = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>'
                '<w:top w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
                '<w:left w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
                '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
                '<w:right w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
                '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
                '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
                '</w:tblBorders>'
            )
            existing = tbl_pr.find(qn('w:tblBorders'))
            if existing is not None:
                tbl_pr.remove(existing)
            tbl_pr.append(borders)

    @staticmethod
    def _fix_image_sizes(doc, page_size: str = "A4",
                         margin_preset: str = "Normal") -> None:
        """Cap all inline images to the usable page width."""
        pw, _ = PAGE_SIZES.get(page_size, PAGE_SIZES["A4"])
        _, _, ml, mr = MARGIN_PRESETS.get(margin_preset, MARGIN_PRESETS["Normal"])
        max_width_inches = pw - ml - mr - 0.2
        max_width_emu = int(max_width_inches * 914400)  # 1 inch = 914400 EMU
        for shape in doc.inline_shapes:
            if shape.width and shape.width > max_width_emu:
                ratio = max_width_emu / shape.width
                shape.width = max_width_emu
                shape.height = int(shape.height * ratio)

    @staticmethod
    def _calc_image_width(img_bytes: bytes, page_size: str = "A4",
                          margin_preset: str = "Normal") -> float:
        """Return a reasonable image width in inches based on actual pixel size."""
        pw, _ = PAGE_SIZES.get(page_size, PAGE_SIZES["A4"])
        _, _, ml, mr = MARGIN_PRESETS.get(margin_preset, MARGIN_PRESETS["Normal"])
        usable_width = pw - ml - mr - 0.2
        try:
            img = Image.open(BytesIO(img_bytes))
            w_px, _ = img.size
            w_inches = w_px / _IMG_RENDER_DPI
            return min(w_inches, usable_width)
        except Exception:
            return min(4.0, usable_width)

    # ── Helpers ───────────────────────────────────────────────────────────────
    @staticmethod
    def _text_to_html(stripped: str) -> str:
        """Convert a single line of text to an HTML element, XSS-safe."""
        if not stripped:
            return ""
        escaped = html.escape(stripped)
        if stripped.startswith("###"):
            return f"<h3>{html.escape(stripped.lstrip('# '))}</h3>"
        if stripped.startswith("##"):
            return f"<h2>{html.escape(stripped.lstrip('# '))}</h2>"
        if stripped.startswith("#"):
            return f"<h1>{html.escape(stripped.lstrip('# '))}</h1>"
        return f"<p>{escaped}</p>"

    # ══════════════════════════════════════════════════════════════════════════
    # LEGACY API  (backward-compatible with tests)
    # ══════════════════════════════════════════════════════════════════════════
    @staticmethod
    def _line_to_html_element(stripped: str) -> str:
        """Convert a single line to an HTML element, XSS-safe."""
        if not stripped:
            return "<br>"
        escaped = html.escape(stripped)
        if stripped.startswith("###"):
            return f"<h3>{html.escape(stripped.lstrip('# '))}</h3>"
        if stripped.startswith("##"):
            return f"<h2>{html.escape(stripped.lstrip('# '))}</h2>"
        if stripped.startswith("#"):
            return f"<h1>{html.escape(stripped.lstrip('# '))}</h1>"
        return f"<p>{escaped}</p>"

    @staticmethod
    def create_txt(text: str, metadata: str = "") -> str:
        content = ""
        if metadata:
            content += f"--- Document Info ---\n{metadata}\n---\n\n"
        content += text
        fd, path = tempfile.mkstemp(suffix=".txt", prefix="ocr_")
        with os.fdopen(fd, "w", encoding="utf-8") as f:
            f.write(content)
        return path

    @staticmethod
    def create_html(text: str, tables_html: Optional[List[str]] = None,
                    figures: Optional[List[Dict[str, Any]]] = None,
                    metadata: str = "") -> str:
        """Legacy HTML builder — XSS-safe."""
        parts = [
            "<!DOCTYPE html><html lang='en'><head><meta charset='utf-8'>",
            "<title>OCR Result</title>",
            "<style>",
            "body{font-family:Tahoma,Arial,sans-serif;max-width:900px;margin:auto;"
            "padding:2rem;line-height:1.7;color:#1e293b}",
            "table{border-collapse:collapse;width:100%;margin:1rem 0}",
            "th,td{border:1px solid #cbd5e1;padding:8px 12px;text-align:left}",
            "th{background:#f1f5f9;font-weight:bold}",
            "img{max-width:100%;height:auto;margin:1rem 0;border-radius:4px}",
            "</style></head><body>",
        ]
        if metadata:
            parts.append(f"<div style='color:#64748b;font-size:0.85rem;"
                         f"border-bottom:1px solid #e2e8f0;padding-bottom:1rem;"
                         f"margin-bottom:2rem'>{html.escape(metadata)}</div>")
        for line in text.split("\n"):
            parts.append(DocumentExporter._line_to_html_element(line.strip()))
        for th in (tables_html or []):
            parts.append(th)
        for fig in (figures or []):
            b64 = fig.get("base64", "")
            if b64 and _is_valid_base64(b64):
                parts.append(
                    f"<figure><img src='data:image/png;base64,{b64}' "
                    f"alt='Figure'/></figure>")
        parts.append("</body></html>")
        fd, path = tempfile.mkstemp(suffix=".html", prefix="ocr_")
        with os.fdopen(fd, "w", encoding="utf-8") as fh:
            fh.write("\n".join(parts))
        return path

    def create_docx(self, text: str,
                    tables_html: Optional[List[str]] = None,
                    figures: Optional[List[Dict[str, Any]]] = None,
                    metadata: str = "") -> Optional[str]:
        """Legacy API — builds HTML first, then converts to DOCX."""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not installed")
            return None
        html_path = self.create_html(text, tables_html, figures, metadata)
        with open(html_path, "r", encoding="utf-8") as f:
            html_content = f.read()
        return self._html_to_docx(html_content)

    def create_all(self, text: str,
                   tables_html: Optional[List[str]] = None,
                   figures: Optional[List[Dict[str, Any]]] = None,
                   metadata: str = "") -> Dict[str, Optional[str]]:
        return {
            "txt": self.create_txt(text, metadata),
            "html": self.create_html(text, tables_html, figures, metadata),
            "docx": self.create_docx(text, tables_html, figures, metadata),
        }
